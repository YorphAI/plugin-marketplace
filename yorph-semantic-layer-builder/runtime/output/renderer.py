"""
Output renderer — converts structured agent outputs into production-ready
semantic layer artifacts.

Always produces TWO files:
  1. The technical format the user chose (dbt YAML, Snowflake, JSON, etc.)
  2. A human-readable Markdown document explaining what was built and why

Design:
  Agent outputs are structured JSON (JOINS_VALIDATED, MEASURES_MB1/2/3, GRAIN_GD1/2/3).
  The renderer takes the chosen recommendation (1=Conservative, 2=Comprehensive, 3=Balanced)
  and the reconciled agent outputs, then generates the appropriate format.

  All business names come from the enriched profiles (documented > humanised).
  Technical names (dbt model names, SQL identifiers) are snake_case.
  Human names (labels, descriptions) are Title Case from documentation.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

import yaml


# ── Canonical intermediate representation ──────────────────────────────────────
# Before writing any format, agents' outputs are normalised into this IR.
# The renderer then serialises IR → target format.

@dataclass
class SemanticEntity:
    """A table/view in the semantic layer (fact or dimension)."""
    technical_name: str          # snake_case, matches warehouse table name
    business_name: str           # human label e.g. "Orders"
    description: str | None
    schema: str
    entity_type: str             # "fact" | "dimension" | "bridge"
    grain: list[str]             # columns that form the grain
    grain_description: str | None
    source_system: str | None    # e.g. "Shopify"
    dimensions: list[str]        # dimension columns (non-measures)
    primary_key: list[str]
    foreign_keys: list["SemanticJoin"]
    extra: dict = field(default_factory=dict)  # pass-through for agent fields not in the schema


@dataclass
class SemanticMeasure:
    """A metric defined on an entity."""
    technical_name: str          # snake_case identifier
    business_name: str           # "Gross Revenue"
    description: str | None
    aggregation: str             # SUM, COUNT, COUNT_DISTINCT, AVG, RATIO
    source_entity: str           # which entity this measure lives on
    source_column: str | None    # the raw column being aggregated
    expression: str | None       # full expression if not a simple column
    filters: list[str]           # WHERE conditions applied before aggregating
    domain: str | None           # "Revenue & Growth", "Customer", etc.
    is_certified: bool = False
    additivity: str = "fully_additive"  # "fully_additive" | "semi_additive" | "non_additive"
    complexity: str = "simple"   # "simple" | "moderate" | "complex"
    numerator: str | None = None # for ratio metrics
    denominator: str | None = None
    time_grains: list[str] = field(default_factory=lambda: ["day", "week", "month", "quarter", "year"])
    notes: str | None = None


@dataclass
class SemanticJoin:
    """A validated join relationship."""
    left_entity: str
    right_entity: str
    left_key: str                # FK column on the left entity
    right_key: str               # PK column on the right entity
    cardinality: str             # "1:1" | "N:1" | "1:many" | "many:many"
    join_type: str = "LEFT"      # "INNER" | "LEFT" | "FULL"
    is_safe: bool = True
    fk_match_rate: float | None = None   # 0.0–1.0, fraction of left rows matching right
    scd_filter: str | None = None        # e.g. "I_REC_END_DATE IS NULL"
    notes: str | None = None
    extra: dict = field(default_factory=dict)  # pass-through for agent fields not in the schema

    @property
    def join_key(self) -> str:
        """Human-readable join key expression for display."""
        if self.left_key and self.right_key:
            return f"{self.left_key} = {self.right_key}"
        return self.left_key or self.right_key or ""

    @property
    def on_clause(self) -> str:
        """SQL ON clause fragment."""
        if self.left_key and self.right_key:
            clause = f"{self.left_entity}.{self.left_key} = {self.right_entity}.{self.right_key}"
            if self.scd_filter:
                clause += f" AND {self.scd_filter}"
            return clause
        return ""


@dataclass
class SemanticMeasure:
    """A metric defined on an entity."""
    technical_name: str          # snake_case identifier
    business_name: str           # "Gross Revenue"
    description: str | None
    aggregation: str             # SUM, COUNT, COUNT_DISTINCT, AVG, RATIO
    source_entity: str           # which entity this measure lives on
    source_column: str | None    # the raw column being aggregated
    expression: str | None       # full expression if not a simple column
    filters: list[str]           # WHERE conditions applied before aggregating
    domain: str | None           # "Revenue & Growth", "Customer", etc.
    is_certified: bool = False
    additivity: str = "fully_additive"  # "fully_additive" | "semi_additive" | "non_additive"
    complexity: str = "simple"   # "simple" | "moderate" | "complex"
    numerator: str | None = None # for ratio metrics
    denominator: str | None = None
    time_grains: list[str] = field(default_factory=lambda: ["day", "week", "month", "quarter", "year"])
    notes: str | None = None
    extra: dict = field(default_factory=dict)  # pass-through for agent fields not in the schema


@dataclass
class SemanticLayer:
    """The complete semantic layer — the IR that everything is rendered from."""
    name: str                    # e.g. "Yorph Semantic Layer — Acme Corp"
    description: str
    warehouse_type: str
    recommendation: str          # "Conservative" | "Comprehensive" | "Balanced"
    generated_at: str
    entities: list[SemanticEntity]
    measures: list[SemanticMeasure]
    joins: list[SemanticJoin]
    business_rules: list[str]    # plain-language rules surfaced during build
    open_questions: list[str]    # unresolved ambiguities to revisit
    glossary: dict[str, str]     # term → definition
    time_intelligence: dict | None = None       # time-based calculations (YoY, MoM, MTD, YTD, etc.)
    dimension_hierarchies: list[dict] | None = None  # drill-down hierarchies (e.g. year → quarter → month)


# ── Format renderers ───────────────────────────────────────────────────────────

class OutputRenderer:
    """
    Renders a SemanticLayer to one of the supported output formats.
    Always also produces a companion markdown document.
    """

    def __init__(self, layer: SemanticLayer, output_dir: Path | None = None):
        self.layer = layer
        self.output_dir = output_dir or (Path.home() / ".yorph" / "output")
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def render(self, fmt: str, filename_base: str | None = None) -> dict[str, Path]:
        """
        Render to the chosen format + always render a companion document.

        Returns dict of {format_key: Path} for every file written.
        """
        base = filename_base or _safe_filename(self.layer.name)
        written = {}

        if fmt == "dbt":
            path = self.output_dir / f"{base}.yml"
            path.write_text(self._render_dbt())
            written["dbt"] = path

        elif fmt == "snowflake":
            path = self.output_dir / f"{base}_snowflake.yml"
            path.write_text(self._render_snowflake_native())
            written["snowflake"] = path

        elif fmt == "json":
            path = self.output_dir / f"{base}.json"
            path.write_text(self._render_json())
            written["json"] = path

        elif fmt == "yaml":
            path = self.output_dir / f"{base}.yaml"
            path.write_text(self._render_yaml())
            written["yaml"] = path

        elif fmt == "osi_spec":
            path = self.output_dir / f"{base}_osi.yaml"
            path.write_text(self._render_osi_spec())
            written["osi_spec"] = path

        elif fmt == "custom":
            path = self.output_dir / f"{base}_custom.yaml"
            path.write_text(self._render_yaml())  # default to YAML for custom
            written["custom"] = path

        elif fmt == "docx":
            path = self.output_dir / f"{base}.docx"
            self._render_docx(path)
            written["docx"] = path

        # Always write the companion document (markdown)
        doc_path = self.output_dir / f"{base}_readme.md"
        doc_path.write_text(self._render_document())
        written["document"] = doc_path

        return written

    # ── dbt format ─────────────────────────────────────────────────────────────

    def _render_dbt(self) -> str:
        """
        Generates dbt-compatible schema.yaml with:
          - models (entities with column descriptions)
          - metrics (dbt Semantic Layer metric definitions)
          - sources (if source system is documented)
        """
        layer = self.layer

        models = []
        for entity in layer.entities:
            columns = []

            # Primary key columns
            for pk in entity.primary_key:
                columns.append({
                    "name": pk,
                    "description": f"Primary key — {entity.grain_description or 'unique identifier'}",
                    "tests": ["unique", "not_null"],
                })

            # Foreign key columns from joins
            for join in layer.joins:
                if join.left_entity == entity.technical_name:
                    fk_col = join.left_key
                    if fk_col and not any(c["name"] == fk_col for c in columns):
                        desc = f"Foreign key → {join.right_entity}.{join.right_key}"
                        if join.scd_filter:
                            desc += f" (SCD filter: {join.scd_filter})"
                        columns.append({
                            "name": fk_col,
                            "description": desc,
                            "tests": ["not_null"] if join.cardinality in ("1:many", "N:1") else [],
                        })

            # Measures as columns
            entity_measures = [m for m in layer.measures if m.source_entity == entity.technical_name]
            for m in entity_measures:
                if m.source_column and not any(c["name"] == m.source_column for c in columns):
                    columns.append({
                        "name": m.source_column,
                        "description": f"{m.business_name} — {m.description or ''}".strip(" —"),
                    })

            model_entry: dict[str, Any] = {
                "name": entity.technical_name,
                "description": entity.description or entity.business_name,
                "columns": [c for c in columns if c],
            }
            if entity.source_system:
                model_entry["meta"] = {"source_system": entity.source_system}
            models.append(model_entry)

        # dbt Semantic Layer metrics
        metrics = []
        for m in layer.measures:
            if m.aggregation == "RATIO":
                # Ratio metrics use derived metric type in dbt
                metric_entry: dict[str, Any] = {
                    "name": m.technical_name,
                    "label": m.business_name,
                    "description": m.description or "",
                    "type": "derived",
                    "type_params": {
                        "expr": m.expression or f"{m.numerator} / NULLIF({m.denominator}, 0)",
                    },
                }
            else:
                metric_entry = {
                    "name": m.technical_name,
                    "label": m.business_name,
                    "description": m.description or "",
                    "type": "simple",
                    "type_params": {
                        "measure": m.technical_name,
                    },
                    "filter": _dbt_filter_string(m.filters) if m.filters else None,
                }
            if m.domain:
                metric_entry["meta"] = {"domain": m.domain}
            metrics.append({k: v for k, v in metric_entry.items() if v is not None})

        # dbt semantic_model entries (MetricFlow format) with entities + measures
        semantic_models = []
        for entity in layer.entities:
            entity_measures = [m for m in layer.measures if m.source_entity == entity.technical_name]
            entity_joins = [j for j in layer.joins if j.left_entity == entity.technical_name and j.is_safe]

            sm_entities = [
                {"name": pk, "type": "primary"} for pk in entity.primary_key
            ]
            for j in entity_joins:
                if j.left_key:
                    sm_entities.append({
                        "name": j.left_key,
                        "type": "foreign",
                        "expr": j.left_key,
                    })

            sm_measures = []
            for m in entity_measures:
                sm_m: dict[str, Any] = {
                    "name": m.technical_name,
                    "agg": m.aggregation.lower(),
                    "description": m.description,
                    "expr": m.expression or m.source_column or m.technical_name,
                }
                if m.filters:
                    sm_m["filter"] = _dbt_filter_string(m.filters)
                sm_measures.append({k: v for k, v in sm_m.items() if v is not None})

            sm_dims = [{"name": d, "type": "categorical"} for d in entity.dimensions[:20]]

            sm: dict[str, Any] = {
                "name": entity.technical_name,
                "description": entity.description or entity.business_name,
                "model": f"ref('{entity.technical_name}')",
                "entities": sm_entities,
                "measures": sm_measures,
                "dimensions": sm_dims,
            }
            semantic_models.append({k: v for k, v in sm.items() if v})

        # Joins section for dbt
        dbt_joins = []
        for j in layer.joins:
            if not j.is_safe:
                continue
            join_entry: dict[str, Any] = {
                "join_to_model": j.right_entity,
                "join_type": j.join_type.lower(),
                "join_on": j.on_clause or None,
                "cardinality": j.cardinality,
            }
            if j.scd_filter:
                join_entry["scd_filter"] = j.scd_filter
            dbt_joins.append({k: v for k, v in join_entry.items() if v is not None})

        doc: dict[str, Any] = {
            "version": 2,
            "models": models,
        }
        if semantic_models:
            doc["semantic_models"] = semantic_models
        if metrics:
            doc["metrics"] = metrics
        if dbt_joins:
            doc["joins"] = dbt_joins

        # Add time intelligence and dimension hierarchies as top-level meta
        meta_block: dict[str, Any] = {}
        if layer.time_intelligence:
            meta_block["time_intelligence"] = layer.time_intelligence
        if layer.dimension_hierarchies:
            meta_block["dimension_hierarchies"] = layer.dimension_hierarchies
        if meta_block:
            doc["meta"] = meta_block

        header = _yaml_comment_header(layer, "dbt Semantic Layer")
        return header + yaml.dump(doc, default_flow_style=False, allow_unicode=True, sort_keys=False)

    # ── Snowflake native semantic views ────────────────────────────────────────

    def _render_snowflake_native(self) -> str:
        """
        Generates Snowflake semantic layer definitions (views + metric DDL style).
        Formatted as YAML that can be applied via Snowflake's semantic layer API
        or manually reviewed before execution.
        """
        layer = self.layer
        entities_out = []

        for entity in layer.entities:
            measures_out = []
            for m in [mm for mm in layer.measures if mm.source_entity == entity.technical_name]:
                expr = m.expression or (
                    f"{m.aggregation}({m.source_column})" if m.source_column
                    else m.aggregation
                )
                if m.filters:
                    filter_str = " AND ".join(m.filters)
                    # Snowflake metric filter uses IFF or WHERE in expression
                    if m.source_column:
                        expr = f"{m.aggregation}(IFF({filter_str}, {m.source_column}, NULL))"

                measure_entry: dict[str, Any] = {
                    "name": m.technical_name,
                    "label": m.business_name,
                    "description": m.description,
                    "expr": expr,
                    "additivity": m.additivity,
                }
                if m.domain:
                    measure_entry["domain"] = m.domain
                measures_out.append({k: v for k, v in measure_entry.items() if v is not None})

            dimensions_out = [{"name": d, "type": "categorical"} for d in entity.dimensions[:20]]

            entity_out: dict[str, Any] = {
                "name": entity.technical_name,
                "label": entity.business_name,
                "description": entity.description,
                "entity_type": entity.entity_type,
                "grain": entity.grain,
                "schema": entity.schema,
                "measures": measures_out,
                "dimensions": dimensions_out,
            }
            if entity.source_system:
                entity_out["source_system"] = entity.source_system
            entities_out.append({k: v for k, v in entity_out.items() if v is not None})

        joins_out = [
            {k: v for k, v in {
                "left": j.left_entity,
                "right": j.right_entity,
                "left_key": j.left_key or None,
                "right_key": j.right_key or None,
                "on": j.on_clause or None,
                "cardinality": j.cardinality,
                "type": j.join_type,
                "fk_match_rate": j.fk_match_rate,
                "scd_filter": j.scd_filter,
                "notes": j.notes,
            }.items() if v is not None}
            for j in layer.joins if j.is_safe
        ]

        doc = {
            "semantic_layer": {
                "name": layer.name,
                "description": layer.description,
                "warehouse": layer.warehouse_type,
                "entities": entities_out,
                "joins": joins_out,
            }
        }

        header = _yaml_comment_header(layer, "Snowflake Semantic Layer")
        return header + yaml.dump(doc, default_flow_style=False, allow_unicode=True, sort_keys=False)

    # ── Generic JSON ──────────────────────────────────────────────────────────

    def _render_json(self) -> str:
        """
        Renders the semantic layer as structured JSON with per-entity column
        classification into facts, dimensions, time_dimensions, and measures.

        Each entity contains:
          - facts:           raw numeric columns referenced by measures
          - dimensions:      categorical / surrogate key columns
          - time_dimensions: date / timestamp columns (with time_grains)
          - measures:        defined aggregation metrics
          - joins:           outbound join relationships for this entity
        """
        layer = self.layer

        entities_out = []
        for entity in layer.entities:
            entity_measures = [m for m in layer.measures
                               if m.source_entity == entity.technical_name]
            entity_joins = [j for j in layer.joins
                            if j.left_entity == entity.technical_name]

            # Collect raw source columns referenced by measures on this entity
            fact_col_names: list[str] = []
            for m in entity_measures:
                if m.source_column and m.source_column not in fact_col_names:
                    fact_col_names.append(m.source_column)

            facts = [
                {
                    "name": col,
                    "expr": col.lower(),
                    "data_type": _infer_data_type(col),
                    "unique": False,
                    "null_ratio": 0,
                    "semantic_type": "numeric",
                }
                for col in fact_col_names
                if not _is_time_column(col)
            ]

            # Dimensions: all columns in entity.dimensions not already in facts/time
            dim_col_names = [
                d for d in entity.dimensions
                if d not in fact_col_names
                and not _is_time_column(d)
            ]
            dimensions = [
                {
                    "name": col,
                    "expr": col.lower(),
                    "data_type": _infer_data_type(col),
                    "unique": False,
                    "null_ratio": 0,
                    "semantic_type": "categorical",
                    "time_grains": [],
                }
                for col in dim_col_names
            ]

            # Time dimensions: date/timestamp columns from grain + dimensions
            time_col_names = [
                d for d in (entity.grain + entity.dimensions)
                if _is_time_column(d) and d not in [t["name"] for t in []]
            ]
            seen_time: set[str] = set()
            time_dimensions = []
            for col in time_col_names:
                if col in seen_time:
                    continue
                seen_time.add(col)
                time_dimensions.append({
                    "name": col,
                    "expr": col.lower(),
                    "data_type": _infer_data_type(col),
                    "unique": False,
                    "null_ratio": 0,
                    "semantic_type": "time",
                    "time_grains": ["day", "week", "month", "quarter", "year"],
                })

            measures_out = [
                {
                    "name": m.technical_name,
                    "label": m.business_name,
                    "description": m.description,
                    "expr": (
                        f"{m.aggregation}({m.source_column})"
                        if m.source_column and m.aggregation not in ("RATIO",)
                        else m.expression or m.source_column
                    ),
                    "aggregation": m.aggregation,
                    "source_column": m.source_column,
                    "additivity": m.additivity,
                    "domain": m.domain,
                    "filters": m.filters or [],
                    "time_grains": m.time_grains,
                    "notes": m.notes,
                }
                for m in entity_measures
            ]

            joins_out = [
                {k: v for k, v in {
                    "right_entity": j.right_entity,
                    "left_key": j.left_key or None,
                    "right_key": j.right_key or None,
                    "on": j.on_clause or None,
                    "cardinality": j.cardinality,
                    "join_type": j.join_type,
                    "safe": j.is_safe,
                    "fk_match_rate": j.fk_match_rate,
                    "scd_filter": j.scd_filter,
                    "notes": j.notes,
                }.items() if v is not None}
                for j in entity_joins
            ]

            entity_out: dict = {
                "name": entity.technical_name,
                "label": entity.business_name,
                "description": entity.description,
                "schema": entity.schema,
                "entity_type": entity.entity_type,
                "grain": entity.grain,
                "grain_description": entity.grain_description,
                "primary_key": entity.primary_key,
                "facts": facts,
                "dimensions": dimensions,
                "time_dimensions": time_dimensions,
                "measures": measures_out,
                "joins": joins_out,
            }
            entities_out.append(
                {k: v for k, v in entity_out.items() if v is not None}
            )

        # Top-level joins section (all joins across all entities)
        all_joins_out = [
            {k: v for k, v in {
                "left_entity": j.left_entity,
                "right_entity": j.right_entity,
                "left_key": j.left_key or None,
                "right_key": j.right_key or None,
                "on": j.on_clause or None,
                "cardinality": j.cardinality,
                "join_type": j.join_type,
                "safe": j.is_safe,
                "fk_match_rate": j.fk_match_rate,
                "scd_filter": j.scd_filter,
                "notes": j.notes,
            }.items() if v is not None}
            for j in layer.joins
        ]

        doc: dict[str, Any] = {
            "semantic_layer": {
                "name": layer.name,
                "description": layer.description,
                "warehouse_type": layer.warehouse_type,
                "recommendation": layer.recommendation,
                "generated_at": layer.generated_at,
                "entities": entities_out,
                "joins": all_joins_out,
                "business_rules": layer.business_rules,
                "open_questions": layer.open_questions,
                "glossary": layer.glossary,
            }
        }
        if layer.time_intelligence:
            doc["semantic_layer"]["time_intelligence"] = layer.time_intelligence
        if layer.dimension_hierarchies:
            doc["semantic_layer"]["dimension_hierarchies"] = layer.dimension_hierarchies
        return json.dumps(doc, indent=2, default=str)

    # ── Generic YAML ──────────────────────────────────────────────────────────

    def _render_yaml(self) -> str:
        header = _yaml_comment_header(self.layer, "Semantic Layer")
        return header + yaml.dump(self._to_dict(), default_flow_style=False, allow_unicode=True, sort_keys=False)

    # ── OSI Spec ──────────────────────────────────────────────────────────────

    def _render_osi_spec(self) -> str:
        """
        Open Semantic Interface (OSI) compatible format.
        Structured for interoperability with semantic layer tools
        that consume a declarative metric spec (e.g. Cube, MetricFlow, Headless BI).
        """
        layer = self.layer

        # Build a lookup: (entity_name, dimension_col) → hierarchy info
        _hierarchy_lookup: dict[tuple[str, str], dict] = {}
        if layer.dimension_hierarchies:
            for h in layer.dimension_hierarchies:
                h_entity = h.get("entity", "")
                h_name = h.get("name", "")
                levels = h.get("levels", [])
                level_names = [
                    (lvl.get("name") if isinstance(lvl, dict) else str(lvl))
                    for lvl in levels
                ]
                for idx, lvl in enumerate(levels):
                    col = lvl.get("name") if isinstance(lvl, dict) else str(lvl)
                    _hierarchy_lookup[(h_entity, col)] = {
                        "hierarchy": h_name,
                        "level_index": idx,
                        "total_levels": len(levels),
                        "levels": level_names,
                    }

        semantic_models = []
        for entity in layer.entities:
            entity_measures = [m for m in layer.measures if m.source_entity == entity.technical_name]
            entity_joins = [j for j in layer.joins
                            if j.left_entity == entity.technical_name and j.is_safe]

            # Build dimensions with optional hierarchy metadata
            dims_out = []
            for d in entity.dimensions[:20]:
                dim_entry: dict[str, Any] = {"name": d, "type": "categorical", "expr": d}
                h_info = _hierarchy_lookup.get((entity.technical_name, d))
                if h_info:
                    dim_entry["hierarchy"] = h_info["hierarchy"]
                    dim_entry["hierarchy_level"] = h_info["level_index"]
                    dim_entry["hierarchy_levels"] = h_info["levels"]
                dims_out.append(dim_entry)

            sm: dict[str, Any] = {
                "name": entity.technical_name,
                "label": entity.business_name,
                "description": entity.description,
                "node_relation": {
                    "schema_name": entity.schema,
                    "alias": entity.technical_name,
                },
                "entities": [
                    {"name": pk, "type": "primary"} for pk in entity.primary_key
                ] + [
                    {k: v for k, v in {
                        "name": j.left_key,
                        "type": "foreign",
                        "expr": j.left_key,
                        "ref_entity": j.right_entity,
                        "ref_key": j.right_key,
                        "scd_filter": j.scd_filter,
                    }.items() if v}
                    for j in entity_joins if j.left_key
                ],
                "measures": [
                    {
                        "name": m.technical_name,
                        "label": m.business_name,
                        "description": m.description,
                        "agg": m.aggregation,
                        "expr": m.source_column or m.expression,
                        "agg_time_dimension": _find_date_dimension(entity),
                    }
                    for m in entity_measures
                    if m.aggregation != "RATIO"
                ],
                "dimensions": dims_out,
            }
            semantic_models.append({k: v for k, v in sm.items() if v is not None})

        metrics = [
            {
                "name": m.technical_name,
                "label": m.business_name,
                "description": m.description,
                "type": "simple" if m.aggregation != "RATIO" else "ratio",
                "type_params": {
                    "measure": m.technical_name,
                    **({"numerator": m.numerator, "denominator": m.denominator}
                       if m.aggregation == "RATIO" else {})
                },
                "filter": _osi_filter_string(m.filters) if m.filters else None,
            }
            for m in layer.measures
        ]

        doc: dict[str, Any] = {
            "version": "1.0",
            "semantic_models": semantic_models,
            "metrics": [{k: v for k, v in m.items() if v is not None} for m in metrics],
        }
        if layer.time_intelligence:
            doc["time_intelligence"] = layer.time_intelligence
        if layer.dimension_hierarchies:
            doc["dimension_hierarchies"] = layer.dimension_hierarchies

        header = _yaml_comment_header(layer, "OSI Semantic Spec")
        return header + yaml.dump(doc, default_flow_style=False, allow_unicode=True, sort_keys=False)

    # ── Microsoft Word (.docx) ─────────────────────────────────────────────────

    def _render_docx(self, path: Path) -> None:
        """
        Generates a Microsoft Word (.docx) document with the same content as
        _render_document() but formatted with proper Word styles: headings,
        tables, bullet lists, bold labels.

        Uses python-docx (already a project dependency).
        Saves directly to `path` — no string intermediate.
        """
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn

        layer = self.layer
        doc = Document()

        # ── Helpers ───────────────────────────────────────────────────────────

        def add_kv_table(rows_data: list[tuple[str, str]]) -> None:
            """Add a 2-column key/value table with bold keys."""
            t = doc.add_table(rows=len(rows_data), cols=2)
            t.style = "Table Grid"
            for i, (label, value) in enumerate(rows_data):
                t.rows[i].cells[0].text = label
                for para in t.rows[i].cells[0].paragraphs:
                    for run in para.runs:
                        run.bold = True
                t.rows[i].cells[1].text = value or "—"

        # ── Title block ───────────────────────────────────────────────────────

        doc.add_heading(layer.name, level=0)

        sub = doc.add_paragraph()
        sub.add_run(f"Generated by Yorph Semantic Layer Assistant · {layer.generated_at[:10]}")
        sub.add_run(f"\nRecommendation: ")
        sub.add_run(layer.recommendation).bold = True
        sub.add_run(f"  |  Warehouse: {layer.warehouse_type}")

        doc.add_paragraph()

        # ── What Was Built ────────────────────────────────────────────────────

        doc.add_heading("What Was Built", level=1)
        doc.add_paragraph(layer.description)
        add_kv_table([
            ("Entities (tables)", str(len(layer.entities))),
            ("Measures (metrics)", str(len(layer.measures))),
            ("Validated joins", str(len(layer.joins))),
            ("Business rules applied", str(len(layer.business_rules))),
        ])
        doc.add_paragraph()

        # ── Entities ──────────────────────────────────────────────────────────

        doc.add_heading("Entities", level=1)

        for entity in layer.entities:
            doc.add_heading(
                f"{entity.business_name}  ({entity.schema}.{entity.technical_name})",
                level=2,
            )

            props: list[tuple[str, str]] = [
                ("Type", entity.entity_type.capitalize()),
                ("Grain", entity.grain_description or ", ".join(entity.grain)),
            ]
            if entity.source_system:
                props.append(("Source system", entity.source_system))
            if entity.primary_key:
                props.append(("Primary key", ", ".join(entity.primary_key)))
            add_kv_table(props)

            if entity.description:
                doc.add_paragraph(entity.description)

            # Joins from this entity
            outbound = [j for j in layer.joins if j.left_entity == entity.technical_name]
            if outbound:
                jp = doc.add_paragraph()
                jp.add_run("Joins:").bold = True
                for j in outbound:
                    safe_label = "safe ✓" if j.is_safe else "use caution ⚠"
                    fk_rate = f"  FK match: {j.fk_match_rate:.0%}" if j.fk_match_rate is not None else ""
                    bullet = doc.add_paragraph(style="List Bullet")
                    bullet.add_run(f"{j.left_key} = {j.right_key}").bold = True
                    bullet.add_run(f"  →  {j.right_entity}  [{j.cardinality}]  ({safe_label}){fk_rate}")
                    if j.scd_filter:
                        scd_note = doc.add_paragraph(style="List Bullet 2")
                        scd_note.add_run(f"SCD filter: {j.scd_filter}").italic = True
                    if j.notes:
                        note = doc.add_paragraph(style="List Bullet 2")
                        note.add_run(j.notes).italic = True

            doc.add_paragraph()

        # ── Metrics ───────────────────────────────────────────────────────────

        doc.add_heading("Metrics", level=1)

        domains: dict[str, list[SemanticMeasure]] = {}
        for m in layer.measures:
            domains.setdefault(m.domain or "General", []).append(m)

        for domain, measures in sorted(domains.items()):
            doc.add_heading(domain, level=2)

            for m in measures:
                cert = "  ✓ certified" if m.is_certified else ""
                doc.add_heading(f"{m.business_name}{cert}", level=3)

                if m.description:
                    doc.add_paragraph(m.description)

                metric_props: list[tuple[str, str]] = [
                    ("Technical name", m.technical_name),
                    ("Formula", m.expression or "—"),
                    ("Source", f"{m.source_entity}.{m.source_column or '—'}"),
                    ("Aggregation", m.aggregation),
                    ("Additivity", m.additivity.replace("_", " ")),
                    ("Complexity", m.complexity),
                ]
                if m.filters:
                    metric_props.append(("Filters", " AND ".join(m.filters)))
                add_kv_table(metric_props)

                if m.notes:
                    note_p = doc.add_paragraph()
                    note_p.add_run("⚠ Note: ").bold = True
                    note_p.add_run(m.notes).italic = True

                doc.add_paragraph()

        # ── Time Intelligence ───────────────────────────────────────────────

        if layer.time_intelligence:
            doc.add_heading("Time Intelligence", level=1)
            ti = layer.time_intelligence

            # Date spine
            date_spine = ti.get("date_spine")
            if isinstance(date_spine, dict):
                if date_spine.get("detected"):
                    props = [
                        ("Table", date_spine.get("table", "—")),
                        ("Date column", date_spine.get("date_column", "—")),
                    ]
                    ds_range = date_spine.get("date_range", {})
                    if ds_range:
                        props.append(("Range", f"{ds_range.get('min', '?')} to {ds_range.get('max', '?')}"))
                    props.append(("Has gaps", "Yes" if date_spine.get("has_gaps") else "No"))
                    doc.add_heading("Date Spine", level=2)
                    add_kv_table(props)
                else:
                    p = doc.add_paragraph()
                    p.add_run("No date spine detected.").bold = True
                    p.add_run(" A date spine is recommended for period-over-period metrics.")
                doc.add_paragraph()

            # Fact time dimensions + calculations
            for ftd in ti.get("fact_time_dimensions", []):
                if not isinstance(ftd, dict):
                    continue
                fact_table = ftd.get("fact_table", "unknown")
                doc.add_heading(fact_table, level=2)
                props = [
                    ("Primary time column", ftd.get("primary_time_column", "—")),
                    ("Grain", ftd.get("grain", "—")),
                ]
                dr = ftd.get("date_range", {})
                if dr:
                    props.append(("Range", f"{dr.get('min', '?')} to {dr.get('max', '?')}"))
                add_kv_table(props)

                calcs = ftd.get("time_calculations", [])
                if calcs:
                    doc.add_paragraph()
                    t = doc.add_table(rows=1, cols=4)
                    t.style = "Table Grid"
                    for i, hdr in enumerate(["Calculation", "Type", "Measure", "Formula"]):
                        t.rows[0].cells[i].text = hdr
                        for para in t.rows[0].cells[i].paragraphs:
                            for run in para.runs:
                                run.bold = True
                    for calc in calcs:
                        if not isinstance(calc, dict):
                            continue
                        row = t.add_row()
                        row.cells[0].text = calc.get("name", "")
                        row.cells[1].text = calc.get("type", "")
                        row.cells[2].text = calc.get("measure", "")
                        row.cells[3].text = calc.get("formula_hint", calc.get("expression", calc.get("formula", "")))
                doc.add_paragraph()

        # ── Business Rules ────────────────────────────────────────────────────

        if layer.business_rules:
            doc.add_heading("Business Rules Applied", level=1)
            doc.add_paragraph(
                "These rules were extracted from documentation and applied to metric definitions:"
            )
            for rule in layer.business_rules:
                doc.add_paragraph(rule, style="List Bullet")
            doc.add_paragraph()

        # ── Open Questions ────────────────────────────────────────────────────

        if layer.open_questions:
            doc.add_heading("⚠ Open Questions to Revisit", level=1)
            doc.add_paragraph(
                "These ambiguities were flagged during the build but not resolved:"
            )
            for q in layer.open_questions:
                doc.add_paragraph(q, style="List Bullet")
            doc.add_paragraph()

        # ── Glossary ──────────────────────────────────────────────────────────

        if layer.glossary:
            doc.add_heading("Glossary", level=1)
            for term, definition in sorted(layer.glossary.items()):
                p = doc.add_paragraph()
                p.add_run(f"{term}: ").bold = True
                p.add_run(definition)

        # ── Footer ────────────────────────────────────────────────────────────

        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(
            f"Generated by Yorph Semantic Layer Assistant · "
            f"{layer.generated_at[:19].replace('T', ' ')} UTC"
        ).italic = True

        doc.save(path)

    # ── Companion document (always generated) ──────────────────────────────────

    def _render_document(self) -> str:
        """
        Human-readable Markdown document — ALWAYS generated alongside the technical format.

        Explains:
        - What was built and why
        - Every entity with its business description
        - Every metric with its formula and business rule
        - All validated joins and their cardinality
        - Open questions the user should revisit
        - Business rules applied
        - Glossary
        """
        layer = self.layer
        lines = [
            f"# {layer.name}",
            f"",
            f"> Generated by Yorph Semantic Layer Assistant on {layer.generated_at[:10]}  ",
            f"> Recommendation: **{layer.recommendation}** | Warehouse: {layer.warehouse_type}",
            f"",
            f"---",
            f"",
            f"## What Was Built",
            f"",
            f"{layer.description}",
            f"",
            f"| | Count |",
            f"|--|--|",
            f"| Entities (tables) | {len(layer.entities)} |",
            f"| Measures (metrics) | {len(layer.measures)} |",
            f"| Validated joins | {len(layer.joins)} |",
            f"| Business rules applied | {len(layer.business_rules)} |",
            f"",
            f"---",
            f"",
            f"## Entities",
            f"",
        ]

        for entity in layer.entities:
            lines += [
                f"### {entity.business_name} (`{entity.schema}.{entity.technical_name}`)",
                f"",
                f"- **Type:** {entity.entity_type.capitalize()}",
                f"- **Grain:** {entity.grain_description or ', '.join(entity.grain)}",
            ]
            if entity.source_system:
                lines.append(f"- **Source system:** {entity.source_system}")
            if entity.description:
                lines += ["", entity.description]
            if entity.primary_key:
                lines.append(f"- **Primary key:** `{'`, `'.join(entity.primary_key)}`")

            # Joins from this entity
            outbound = [j for j in layer.joins if j.left_entity == entity.technical_name]
            if outbound:
                lines += ["", "**Joins:**"]
                for j in outbound:
                    safe = "✅ safe" if j.is_safe else "⚠ use caution"
                    fk_rate = f" (FK match: {j.fk_match_rate:.0%})" if j.fk_match_rate is not None else ""
                    lines.append(
                        f"- `{j.left_key}` = `{j.right_entity}.{j.right_key}` [{j.cardinality}] {safe}{fk_rate}"
                    )
                    if j.scd_filter:
                        lines.append(f"  - SCD filter: `{j.scd_filter}`")
                    if j.notes:
                        lines.append(f"  - {j.notes}")

            lines.append("")

        lines += [
            f"---",
            f"",
            f"## Metrics",
            f"",
        ]

        # Group by domain
        domains: dict[str, list[SemanticMeasure]] = {}
        for m in layer.measures:
            domains.setdefault(m.domain or "General", []).append(m)

        for domain, measures in sorted(domains.items()):
            lines += [f"### {domain}", ""]
            for m in measures:
                cert = " ✓ *certified*" if m.is_certified else ""
                lines += [f"#### {m.business_name}{cert}", ""]
                if m.description:
                    lines.append(m.description)
                lines += [
                    f"",
                    f"| Property | Value |",
                    f"|--|--|",
                    f"| Technical name | `{m.technical_name}` |",
                    f"| Formula | `{m.expression or '—'}` |",
                    f"| Source | `{m.source_entity}.{m.source_column or '—'}` |",
                    f"| Aggregation | {m.aggregation} |",
                    f"| Additivity | {m.additivity.replace('_', ' ')} |",
                    f"| Complexity | {m.complexity} |",
                ]
                if m.filters:
                    lines.append(f"| Filters | `{' AND '.join(m.filters)}` |")
                if m.notes:
                    lines += ["", f"> ⚠ {m.notes}"]
                lines.append("")

        if layer.business_rules:
            lines += [
                "---",
                "",
                "## Business Rules Applied",
                "",
                "These rules were extracted from documentation and applied to metric definitions:",
                "",
            ]
            for rule in layer.business_rules:
                lines.append(f"- {rule}")
            lines.append("")

        if layer.open_questions:
            lines += [
                "---",
                "",
                "## ⚠ Open Questions to Revisit",
                "",
                "These ambiguities were flagged during the build but not resolved:",
                "",
            ]
            for q in layer.open_questions:
                lines.append(f"- {q}")
            lines.append("")

        if layer.glossary:
            lines += [
                "---",
                "",
                "## Glossary",
                "",
            ]
            for term, definition in sorted(layer.glossary.items()):
                lines.append(f"**{term}**: {definition}")
                lines.append("")

        if layer.time_intelligence:
            lines += [
                "---",
                "",
                "## Time Intelligence",
                "",
                "Pre-defined time-based calculations available for measures:",
                "",
            ]
            ti = layer.time_intelligence

            # Date spine info
            date_spine = ti.get("date_spine")
            if isinstance(date_spine, dict):
                detected = date_spine.get("detected", False)
                if detected:
                    ds_table = date_spine.get("table", "unknown")
                    ds_col = date_spine.get("date_column", "")
                    ds_range = date_spine.get("date_range", {})
                    lines.append(f"### Date Spine: `{ds_table}`")
                    lines.append("")
                    lines.append(f"- **Date column:** `{ds_col}`")
                    if ds_range:
                        lines.append(f"- **Range:** {ds_range.get('min', '?')} to {ds_range.get('max', '?')}")
                    lines.append(f"- **Has gaps:** {'Yes' if date_spine.get('has_gaps') else 'No'}")
                    lines.append("")
                else:
                    lines.append("> **No date spine detected.** A date spine is recommended for period-over-period metrics.")
                    lines.append("")

            # Fact time dimensions + calculations
            fact_dims = ti.get("fact_time_dimensions", [])
            for ftd in fact_dims:
                if not isinstance(ftd, dict):
                    continue
                fact_table = ftd.get("fact_table", "unknown")
                time_col = ftd.get("primary_time_column", "")
                grain_val = ftd.get("grain", "")
                lines.append(f"### {fact_table}")
                lines.append("")
                lines.append(f"- **Primary time column:** `{time_col}` ({grain_val} grain)")
                date_range = ftd.get("date_range", {})
                if date_range:
                    lines.append(f"- **Range:** {date_range.get('min', '?')} to {date_range.get('max', '?')}")
                lines.append("")

                calcs = ftd.get("time_calculations", [])
                if calcs:
                    lines.append("| Calculation | Type | Measure | Formula |")
                    lines.append("|--|--|--|--|")
                    for calc in calcs:
                        if not isinstance(calc, dict):
                            continue
                        c_name = calc.get("name", "")
                        c_type = calc.get("type", "")
                        c_measure = calc.get("measure", "")
                        c_formula = calc.get("formula_hint", calc.get("expression", calc.get("formula", "")))
                        lines.append(f"| `{c_name}` | {c_type} | `{c_measure}` | `{c_formula}` |")
                    lines.append("")

            # Fiscal calendar
            fiscal = ti.get("fiscal_calendar")
            if fiscal and isinstance(fiscal, dict):
                lines.append("### Fiscal Calendar")
                lines.append("")
                for fk, fv in fiscal.items():
                    lines.append(f"- **{fk}:** {fv}")
                lines.append("")

            # Fallback: render any top-level keys not already handled
            handled_keys = {"date_spine", "fact_time_dimensions", "fiscal_calendar", "assumption_questions"}
            for calc_name, calc_def in ti.items():
                if calc_name in handled_keys:
                    continue
                if isinstance(calc_def, dict):
                    label = calc_def.get("label", calc_name)
                    desc = calc_def.get("description", "")
                    expr = calc_def.get("expression", calc_def.get("formula_hint", calc_def.get("formula", "")))
                    lines.append(f"### {label}")
                    lines.append("")
                    if desc:
                        lines.append(desc)
                        lines.append("")
                    if expr:
                        lines.append(f"- **Expression:** `{expr}`")
                    applicable = calc_def.get("applicable_measures")
                    if applicable:
                        measures_str = ", ".join(f"`{m}`" for m in applicable)
                        lines.append(f"- **Applicable measures:** {measures_str}")
                    grain_val = calc_def.get("grain")
                    if grain_val:
                        lines.append(f"- **Grain:** {grain_val}")
                    lines.append("")
                else:
                    lines.append(f"- **{calc_name}**: {calc_def}")
                    lines.append("")

        if layer.dimension_hierarchies:
            lines += [
                "---",
                "",
                "## Dimension Hierarchies",
                "",
                "Drill-down paths defined for dimensional analysis:",
                "",
            ]
            for hierarchy in layer.dimension_hierarchies:
                h_name = hierarchy.get("name", "Unnamed Hierarchy")
                h_entity = hierarchy.get("entity", "")
                h_levels = hierarchy.get("levels", [])
                lines.append(f"### {h_name}")
                lines.append("")
                if h_entity:
                    lines.append(f"- **Entity:** `{h_entity}`")
                if h_levels:
                    level_strs = []
                    for lvl in h_levels:
                        if isinstance(lvl, dict):
                            level_strs.append(lvl.get("name", str(lvl)))
                        else:
                            level_strs.append(str(lvl))
                    lines.append(f"- **Levels:** {' -> '.join(level_strs)}")
                desc = hierarchy.get("description")
                if desc:
                    lines.append(f"- **Description:** {desc}")
                lines.append("")

        lines += [
            "---",
            "",
            f"*Generated by Yorph Semantic Layer Assistant · {layer.generated_at[:19].replace('T', ' ')} UTC*",
        ]

        return "\n".join(lines)

    # ── Shared serialisation ──────────────────────────────────────────────────

    def _to_dict(self) -> dict:
        layer = self.layer
        sl: dict[str, Any] = {
            "name": layer.name,
            "description": layer.description,
            "warehouse_type": layer.warehouse_type,
            "recommendation": layer.recommendation,
            "generated_at": layer.generated_at,
            "entities": [
                {
                    "technical_name": e.technical_name,
                    "business_name": e.business_name,
                    "description": e.description,
                    "schema": e.schema,
                    "entity_type": e.entity_type,
                    "grain": e.grain,
                    "grain_description": e.grain_description,
                    "source_system": e.source_system,
                    "primary_key": e.primary_key,
                    "dimensions": e.dimensions,
                }
                for e in layer.entities
            ],
            "measures": [
                {k: v for k, v in {
                    "technical_name": m.technical_name,
                    "business_name": m.business_name,
                    "description": m.description,
                    "expr": m.expression,
                    "aggregation": m.aggregation,
                    "source_entity": m.source_entity,
                    "source_column": m.source_column,
                    "filters": m.filters if m.filters else None,
                    "domain": m.domain,
                    "is_certified": m.is_certified if m.is_certified else None,
                    "additivity": m.additivity,
                    "complexity": m.complexity,
                    "time_grains": m.time_grains,
                    "numerator": m.numerator,
                    "denominator": m.denominator,
                    "notes": m.notes,
                }.items() if v is not None}
                for m in layer.measures
            ],
            "joins": [
                {k: v for k, v in {
                    "left_entity": j.left_entity,
                    "right_entity": j.right_entity,
                    "left_key": j.left_key or None,
                    "right_key": j.right_key or None,
                    "on": j.on_clause or None,
                    "cardinality": j.cardinality,
                    "join_type": j.join_type,
                    "is_safe": j.is_safe,
                    "fk_match_rate": j.fk_match_rate,
                    "scd_filter": j.scd_filter,
                    "notes": j.notes,
                }.items() if v is not None}
                for j in layer.joins
            ],
            "business_rules": layer.business_rules,
            "open_questions": layer.open_questions,
            "glossary": layer.glossary,
        }
        if layer.time_intelligence:
            sl["time_intelligence"] = layer.time_intelligence
        if layer.dimension_hierarchies:
            sl["dimension_hierarchies"] = layer.dimension_hierarchies
        return {"semantic_layer": sl}


# ── Helpers ────────────────────────────────────────────────────────────────────

def _safe_filename(name: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_\-]", "_", name.lower()).strip("_")[:60]


def _yaml_comment_header(layer: SemanticLayer, format_name: str) -> str:
    return (
        f"# {format_name}\n"
        f"# Generated by Yorph Semantic Layer Assistant\n"
        f"# Recommendation: {layer.recommendation}\n"
        f"# Generated: {layer.generated_at[:19].replace('T', ' ')} UTC\n"
        f"# Warehouse: {layer.warehouse_type}\n"
        f"#\n"
        f"# ⚠ Review before applying to production.\n"
        f"# See the companion _readme.md for full explanation of every decision.\n\n"
    )


def _dbt_filter_string(filters: list[str]) -> str:
    """Convert filters list to dbt Jinja filter string."""
    conditions = " and ".join(f"{{{{ Dimension('{f.split('=')[0].strip()}') }}}} = '{f.split('=')[-1].strip().strip(chr(39))}'" for f in filters)
    return f"{{{{{{ where: \"{conditions}\" }}}}}}"


def _osi_filter_string(filters: list[str]) -> str:
    return "{{{{ where: \"" + " AND ".join(filters) + "\" }}}}"


def _is_time_column(col: str) -> bool:
    """Return True if the column name looks like a date or timestamp."""
    lower = col.lower()
    time_suffixes = ("_date_sk", "_time_sk", "_date", "_at", "_ts", "_timestamp",
                     "_month", "_year", "_day", "_week", "_quarter")
    time_substrings = ("date", "time", "month", "year", "quarter", "week", "day")
    return (
        any(lower.endswith(s) for s in time_suffixes)
        or any(s in lower for s in time_substrings)
    )


def _infer_data_type(col: str) -> str:
    """Infer a SQL data type string from column naming conventions."""
    lower = col.lower()
    if any(lower.endswith(s) for s in ("_date", "_at", "_date_sk")):
        return "DATE"
    if any(lower.endswith(s) for s in ("_timestamp", "_ts", "_time_sk")):
        return "TIMESTAMP"
    if any(lower.endswith(s) for s in ("_month", "_year", "_day", "_week", "_quarter")):
        return "DATE"
    if any(lower.endswith(s) for s in ("_sk", "_key", "_id", "_number", "_count",
                                        "_qty", "_quantity", "_amount", "_cost",
                                        "_price", "_paid", "_profit", "_tax",
                                        "_fraction", "_rate", "_score", "_pct")):
        return "NUMBER"
    if any(s in lower for s in ("name", "type", "status", "flag", "code",
                                 "category", "label", "color", "group", "url",
                                 "pokemon", "brand")):
        return "STRING"
    return "NUMBER"  # safe default for warehouse numeric columns


def _find_date_dimension(entity: SemanticEntity) -> str | None:
    """Find the primary date/timestamp dimension for an entity."""
    date_hints = ["created_at", "order_date", "event_date", "date", "timestamp",
                  "created_date", "updated_at", "transaction_date"]
    for hint in date_hints:
        if hint in [d.lower() for d in entity.dimensions]:
            return hint
    for d in entity.dimensions:
        if any(t in d.lower() for t in ["date", "time", "at", "ts"]):
            return d
    return None


# ── Factory: build SemanticLayer from agent outputs ───────────────────────────

def build_semantic_layer_from_agent_outputs(
    agent_outputs: dict,
    recommendation_number: int = 3,
    warehouse_type: str = "unknown",
    project_name: str = "Semantic Layer",
    description: str = "",
    joins_grade: int | None = None,
    measures_grade: int | None = None,
    grain_grade: int | None = None,
) -> SemanticLayer:
    """
    Parses structured agent JSON outputs and constructs a SemanticLayer IR.

    Supports both bundled recommendations (recommendation_number) and
    independent per-dimension grade selection (joins_grade, measures_grade,
    grain_grade). Individual grade params take priority over recommendation_number.

    Args:
        agent_outputs: dict with keys:
            "joins_jv1" / "joins_jv2" / "joins_jv3" (or fallback "joins"),
            "measures_mb1", "measures_mb2", "measures_mb3",
            "grain_gd1", "grain_gd2", "grain_gd3",
            "business_rules", "open_questions", "glossary"
        recommendation_number: 1=Conservative, 2=Comprehensive, 3=Balanced.
            Used as the default for any dimension grade not individually specified.
        joins_grade: Override joins tier independently (1=Strict, 2=Explorer, 3=Trap Hunter)
        measures_grade: Override measures tier independently (1=Minimalist, 2=Analyst, 3=Strategist)
        grain_grade: Override grain tier independently (1=Purist, 2=Pragmatist, 3=Architect)
    """
    grade_names = {1: "Conservative", 2: "Comprehensive", 3: "Balanced"}

    # Resolve effective grade per dimension — individual override > recommendation_number
    jg = joins_grade if joins_grade in (1, 2, 3) else recommendation_number
    mg = measures_grade if measures_grade in (1, 2, 3) else recommendation_number
    gg = grain_grade if grain_grade in (1, 2, 3) else recommendation_number

    # Grade → label per dimension
    joins_labels  = {1: "JV-1 Strict", 2: "JV-2 Explorer", 3: "JV-3 Trap Hunter"}
    measures_labels = {1: "MB-1 Minimalist", 2: "MB-2 Analyst", 3: "MB-3 Strategist"}
    grain_labels  = {1: "GD-1 Purist", 2: "GD-2 Pragmatist", 3: "GD-3 Architect"}

    jg_label = joins_labels.get(jg, "JV-3 Trap Hunter")
    mg_label = measures_labels.get(mg, "MB-3 Strategist")
    gg_label = grain_labels.get(gg, "GD-3 Architect")

    # Build a human-readable recommendation name
    all_same = (jg == mg == gg)
    if all_same:
        rec_name = grade_names.get(jg, "Balanced")
    else:
        rec_name = f"Custom ({jg_label} / {mg_label} / {gg_label})"

    # Resolve joins: prefer per-grade key, fall back to the single "joins" list
    joins_raw_key = f"joins_jv{jg}"
    raw_joins = agent_outputs.get(joins_raw_key) or agent_outputs.get("joins", [])
    joins = _parse_joins(raw_joins)

    # Resolve measures and grain
    measures_key = f"measures_mb{mg}"
    grain_key = f"grain_gd{gg}"

    raw_measures = agent_outputs.get(measures_key, [])
    measures = _parse_measures(raw_measures)

    raw_grain = agent_outputs.get(grain_key, [])
    domain_context = agent_outputs.get("domain_context", {})
    entities = _parse_entities(raw_grain, warehouse_type, joins, measures, domain_context)

    return SemanticLayer(
        name=f"{project_name} — {rec_name}",
        description=description or f"{rec_name} semantic layer design.",
        warehouse_type=warehouse_type,
        recommendation=rec_name,
        generated_at=datetime.utcnow().isoformat(),
        entities=entities,
        measures=measures,
        joins=joins,
        business_rules=agent_outputs.get("business_rules", []),
        open_questions=agent_outputs.get("open_questions", []),
        glossary=agent_outputs.get("glossary", {}),
        time_intelligence=agent_outputs.get("time_intelligence"),
        dimension_hierarchies=agent_outputs.get("dimension_hierarchies"),
    )


def _parse_joins(raw: list[dict]) -> list[SemanticJoin]:
    joins = []
    for r in raw:
        left, right = _parse_join_tables(r.get("join", ""))
        # Accept both "left_table"/"right_table" and "left_entity"/"right_entity"
        left_entity = left or r.get("left_table", r.get("left_entity", ""))
        right_entity = right or r.get("right_table", r.get("right_entity", ""))
        # Accept "left_key"/"right_key" or fall back to legacy "join_key"
        left_key = r.get("left_key", "")
        right_key = r.get("right_key", "")
        if not left_key and not right_key:
            # Legacy format: single join_key string
            legacy_key = r.get("join_key", "")
            if "=" in legacy_key:
                parts = [p.strip() for p in legacy_key.split("=")]
                left_key, right_key = parts[0], parts[1] if len(parts) > 1 else ""
            else:
                left_key = legacy_key
                right_key = legacy_key
        # Accept "fan_out_safe" as alias for "safe"
        is_safe = r.get("safe", r.get("fan_out_safe", True))
        # FK match rate: accept float or percentage string
        fk_rate = r.get("fk_match_rate")
        if isinstance(fk_rate, str):
            fk_rate = float(fk_rate.rstrip("%")) / (100 if float(fk_rate.rstrip("%")) > 1 else 1)

        joins.append(SemanticJoin(
            left_entity=left_entity,
            right_entity=right_entity,
            left_key=left_key,
            right_key=right_key,
            cardinality=r.get("cardinality", "N:1"),
            is_safe=is_safe,
            fk_match_rate=fk_rate,
            scd_filter=r.get("scd_filter"),
            notes=r.get("notes"),
        ))
    return joins


def _parse_join_tables(join_str: str) -> tuple[str, str]:
    """Parse 'orders → order_items' into ('orders', 'order_items')."""
    if "→" in join_str:
        parts = [p.strip() for p in join_str.split("→")]
        return parts[0], parts[1] if len(parts) > 1 else ("", "")
    return "", ""


def _parse_measures(raw: list[dict]) -> list[SemanticMeasure]:
    measures = []
    for r in raw:
        # Accept multiple field name conventions from agent outputs
        name = r.get("name", r.get("label", r.get("measure_id", "measure")))
        tech_name = r.get("measure_id", _safe_filename(name))
        business_name = r.get("label", name)

        # Source entity: accept "table", "source_table", or "source_entity"
        source_entity = r.get("source_table", r.get("table", r.get("source_entity", "")))
        # For cross-channel measures, "tables" (plural) may be a list
        if not source_entity and "tables" in r:
            tables = r["tables"]
            source_entity = tables[0] if isinstance(tables, list) and tables else ""

        # Parse aggregation + source_column from formula if not provided separately
        aggregation = r.get("aggregation", "")
        source_column = r.get("source_column")
        formula = r.get("formula", r.get("expression", ""))

        if not aggregation and formula:
            # Extract aggregation from formula like "SUM(SS_NET_PAID)"
            agg_match = re.match(r"^(SUM|COUNT|AVG|MIN|MAX|COUNT_DISTINCT)\((.+)\)$",
                                 formula.strip(), re.IGNORECASE)
            if agg_match:
                aggregation = agg_match.group(1).upper()
                if not source_column:
                    inner = agg_match.group(2).strip()
                    # Only set source_column if it's a simple column ref (no operators)
                    if re.match(r"^[A-Za-z_][A-Za-z0-9_]*$", inner):
                        source_column = inner
            elif "COUNT(DISTINCT" in formula.upper():
                aggregation = "COUNT_DISTINCT"
                col_match = re.search(r"COUNT\(DISTINCT\s+(\w+)\)", formula, re.IGNORECASE)
                if col_match and not source_column:
                    source_column = col_match.group(1)
            elif "/" in formula or "+" in formula or "-" in formula:
                aggregation = "RATIO" if "/" in formula else "SUM"
            else:
                aggregation = "SUM"
        if not aggregation:
            aggregation = "SUM"

        # Map "type" field to additivity
        measure_type = r.get("type", "additive")
        additivity_map = {
            "additive": "fully_additive",
            "cross_channel_additive": "fully_additive",
            "derived_ratio": "non_additive",
            "ratio": "non_additive",
            "semi_additive": "semi_additive",
        }
        additivity = r.get("additivity", additivity_map.get(measure_type, "fully_additive"))

        # Complexity: ratio metrics are moderate, cross-channel are moderate
        complexity = r.get("complexity", "simple")
        if aggregation == "RATIO" or measure_type == "derived_ratio":
            complexity = "moderate"
        elif measure_type == "cross_channel_additive":
            complexity = "moderate"

        # Filters
        raw_filter = r.get("filter", r.get("filters", ""))
        if isinstance(raw_filter, list):
            filters = raw_filter
        elif isinstance(raw_filter, str) and raw_filter:
            filters = raw_filter.split(" AND ")
        else:
            filters = []

        # Detect ratio measures that the agent mislabeled as SUM
        # If the name contains "ratio", "pct", "percent", "rate", "fraction", "margin"
        # and there's no source_column and no formula, it's likely a derived ratio
        # that shouldn't be SUM-aggregated
        ratio_signals = ("ratio", "pct", "percent", "rate", "fraction", "margin", "yield")
        name_lower = tech_name.lower()
        if (aggregation == "SUM"
                and not source_column
                and not formula
                and any(sig in name_lower for sig in ratio_signals)):
            aggregation = "RATIO"
            additivity = "non_additive"
            complexity = "moderate"
            # Try to infer numerator/denominator from the name
            # e.g., "cost_to_revenue_ratio" → numerator=cost, denominator=revenue
            if not r.get("numerator") and "_to_" in name_lower:
                parts = name_lower.split("_to_")
                if len(parts) == 2:
                    num_part = parts[0].replace("_", " ").strip()
                    den_part = parts[1].replace("_ratio", "").replace("_pct", "").replace("_", " ").strip()
                    # Store as hints — the agent should provide actual formulas
                    formula = f"SUM({num_part}) / NULLIF(SUM({den_part}), 0)"

        # Always compute the full SQL expression so every output format has it
        if formula:
            # Agent provided an explicit formula — use it as-is
            full_expr = formula
        elif source_column and aggregation:
            if filters:
                filter_str = " AND ".join(filters)
                full_expr = f"{aggregation}(IFF({filter_str}, {source_column}, NULL))"
            else:
                full_expr = f"{aggregation}({source_column})"
        else:
            full_expr = aggregation or ""

        measures.append(SemanticMeasure(
            technical_name=tech_name,
            business_name=business_name,
            description=r.get("description"),
            aggregation=aggregation,
            source_entity=source_entity,
            source_column=source_column,
            expression=full_expr or None,
            filters=filters,
            domain=r.get("domain"),
            is_certified=r.get("is_certified", False),
            additivity=additivity,
            complexity=complexity,
            numerator=r.get("numerator"),
            denominator=r.get("denominator"),
            notes=r.get("notes"),
        ))
    return measures


def _parse_entities(raw_grain: list[dict], warehouse_type: str,
                    joins: list[SemanticJoin],
                    measures: list[SemanticMeasure],
                    domain_context: dict | None = None) -> list[SemanticEntity]:
    entities = []
    measure_tables = {m.source_entity for m in measures}
    join_tables = {j.left_entity for j in joins} | {j.right_entity for j in joins}
    ctx = domain_context or {}

    for r in raw_grain:
        table_name = r.get("table", "")
        # Accept "grain_columns" (from agent output) or "grain"/"reporting_grain"
        raw_grain_val = r.get("grain_columns", r.get("grain", r.get("reporting_grain", [])))
        if isinstance(raw_grain_val, str):
            raw_grain_val = [raw_grain_val]

        # Separate actual column names from description strings.
        # Column names look like: "user_id", "SS_TICKET_NUMBER", "order_id"
        # Descriptions look like: "One row per store ticket × item (atomic)"
        grain_cols: list[str] = []
        grain_desc_parts: list[str] = []
        for item in raw_grain_val:
            # A column name is a single identifier (letters, digits, underscores only)
            if re.match(r'^[A-Za-z_][A-Za-z0-9_]*$', item.strip()):
                grain_cols.append(item.strip())
            else:
                grain_desc_parts.append(item)

        # Look up domain_context for this table
        table_ctx = ctx.get(table_name, {})
        annotated_cols = table_ctx.get("annotated_columns", {})

        # Determine entity type: prefer domain_context, then grain data, then heuristic
        entity_type = r.get("grain_type", table_ctx.get("table_type", ""))
        if not entity_type:
            if table_name in measure_tables:
                entity_type = "fact"
            elif table_name not in join_tables:
                entity_type = "dimension"
            else:
                entity_type = "fact"

        # Description: prefer domain_context, then grain data, then parsed grain descriptions
        description = (
            table_ctx.get("description")
            or r.get("grain_description", r.get("description"))
            or ("; ".join(grain_desc_parts) if grain_desc_parts else None)
        )

        # Grain columns: prefer parsed columns, then domain_context, then primary_key from agent
        if not grain_cols and table_ctx.get("grain"):
            grain_cols = table_ctx["grain"]
        if not grain_cols:
            # Try to extract from primary_key field in agent output
            pk_val = r.get("primary_key", [])
            if isinstance(pk_val, str):
                pk_val = [pk_val]
            grain_cols = [c for c in pk_val if re.match(r'^[A-Za-z_][A-Za-z0-9_]*$', c.strip())]

        # Grain description: use explicit field, parsed descriptions, or build from columns
        grain_description = (
            r.get("grain_description", r.get("notes"))
            or ("; ".join(grain_desc_parts) if grain_desc_parts else None)
        )

        # Primary key = grain columns (they're the same thing in a semantic layer)
        primary_key = grain_cols

        # Classify annotated columns into dimensions (FKs) and measures
        dimension_cols = []
        for col_name, col_info in annotated_cols.items():
            role = col_info.get("role", "") if isinstance(col_info, dict) else ""
            if role.endswith("_fk") or role == "order_id":
                dimension_cols.append(col_name)

        # Also include dimensions from raw grain data
        if not dimension_cols:
            dimension_cols = r.get("safe_dimensions", [])

        # Entity name from grain data, or derive from table name
        entity_name = r.get("entity", "")
        business_name = r.get("business_name",
                              _safe_filename(entity_name or table_name).replace("_", " ").title())

        entities.append(SemanticEntity(
            technical_name=table_name,
            business_name=business_name,
            description=description,
            schema=r.get("schema", table_name.split(".")[0] if "." in table_name else "PUBLIC"),
            entity_type=entity_type,
            grain=grain_cols,
            grain_description=grain_description,
            source_system=r.get("source_system"),
            dimensions=dimension_cols,
            primary_key=primary_key,
            foreign_keys=[j for j in joins if j.left_entity == table_name],
        ))

    return entities
