"""
Document processor — extracts structured DocumentContext from uploaded files and URLs.

Supported inputs:
  Files:  .pdf, .docx, .csv, .json, .yaml/.yml, .xlsx, .md, .txt
  URLs:   any HTTP/HTTPS URL (Confluence, Notion public, GitHub, internal wikis)

Extraction strategy:
  1. Parse the raw content (format-specific)
  2. Run Claude-guided extraction to identify tables, columns, metrics, rules, glossary
  3. Return a structured DocumentContext (never raw text in context)

The extraction is imperfect — confidence ratings signal how reliable the output is.
Agents are instructed to treat lower-confidence extractions as hints, not facts.
"""

from __future__ import annotations

import csv
import io
import json
import re
from pathlib import Path
from typing import Any

import httpx
import yaml

from .context import (
    DocumentContext, TableDefinition, ColumnDefinition,
    MeasureDefinition, MetricDefinition, BusinessRule, GlossaryTerm, JoinHint,
    save_document_context,
)


# ── File type routing ──────────────────────────────────────────────────────────

EXTENSION_MAP = {
    ".pdf":  "pdf",
    ".docx": "docx",
    ".doc":  "docx",
    ".xlsx": "xlsx",
    ".xls":  "xlsx",
    ".csv":  "csv",
    ".json": "json",
    ".yaml": "yaml",
    ".yml":  "yaml",
    ".md":   "markdown",
    ".txt":  "text",
}


def process_file(file_path: str, document_type: str = "unknown") -> DocumentContext:
    """
    Parse a local file into a DocumentContext.

    Args:
        file_path:     absolute path to the file
        document_type: user-specified type hint ("data_dictionary", "saas_context",
                       "business_glossary", "existing_semantic_layer", "schema_docs")
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    source_type = EXTENSION_MAP.get(path.suffix.lower(), "text")

    raw_text = _extract_text(path, source_type)
    ctx = _parse_into_context(
        raw_text=raw_text,
        source_path=file_path,
        source_type=source_type,
        document_type=document_type,
    )
    save_document_context(ctx)
    return ctx


def fetch_url(url: str, document_type: str = "schema_docs") -> DocumentContext:
    """
    Fetch a URL and extract DocumentContext from the page content.

    Handles:
    - Confluence pages (strips navigation chrome)
    - GitHub README / wiki pages
    - Notion public pages
    - Generic HTML pages (best-effort)
    - Raw JSON/YAML endpoints (e.g. dbt manifest.json hosted on S3)
    """
    raw_text, detected_source_type = _fetch_url_content(url)
    ctx = _parse_into_context(
        raw_text=raw_text,
        source_path=url,
        source_type=detected_source_type,
        document_type=document_type,
    )
    save_document_context(ctx)
    return ctx


# ── Raw text extraction (per file format) ─────────────────────────────────────

def _extract_text(path: Path, source_type: str) -> str:
    """Extract raw text content from a file. Format-specific."""

    if source_type == "pdf":
        return _extract_pdf(path)

    elif source_type == "docx":
        return _extract_docx(path)

    elif source_type == "xlsx":
        return _extract_xlsx(path)

    elif source_type == "csv":
        return _extract_csv(path)

    elif source_type == "json":
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            data = json.load(f)
        return json.dumps(data, indent=2)[:50_000]  # cap at 50K chars

    elif source_type == "yaml":
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            data = yaml.safe_load(f)
        return yaml.dump(data)[:50_000]

    else:  # markdown, text
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()[:50_000]


def _extract_pdf(path: Path) -> str:
    """Extract text from PDF using pdfplumber (preserves table structure)."""
    try:
        import pdfplumber
        pages_text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                # Try to extract tables first (better structure than raw text)
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        rows = ["\t".join(str(cell or "") for cell in row) for row in table]
                        pages_text.append("\n".join(rows))
                else:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
        return "\n\n".join(pages_text)[:80_000]
    except ImportError:
        raise ImportError(
            "pdfplumber is required for PDF processing. "
            "Run: pip install pdfplumber"
        )


def _extract_docx(path: Path) -> str:
    """Extract text from Word document."""
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                parts.append(row_text)
        return "\n".join(parts)[:80_000]
    except ImportError:
        raise ImportError(
            "python-docx is required for Word document processing. "
            "Run: pip install python-docx"
        )


def _extract_xlsx(path: Path) -> str:
    """Extract text from Excel file, sheet by sheet."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)[:80_000]
    except ImportError:
        raise ImportError(
            "openpyxl is required for Excel processing. "
            "Run: pip install openpyxl"
        )


def _extract_csv(path: Path) -> str:
    """Return raw CSV content (preserving original delimiters for DictReader)."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()[:80_000]


# ── URL fetching ───────────────────────────────────────────────────────────────

def _fetch_url_content(url: str) -> tuple[str, str]:
    """
    Fetch a URL and return (raw_text, detected_source_type).

    Handles:
    - Raw JSON/YAML endpoints → structured parse
    - GitHub raw content → text
    - HTML pages → strip to main content via BeautifulSoup
    """
    headers = {
        "User-Agent": "Yorph-SemanticLayer-Agent/0.1 (documentation reader)",
        "Accept": "text/html,application/xhtml+xml,application/json,text/plain,*/*",
    }

    try:
        response = httpx.get(url, headers=headers, follow_redirects=True, timeout=30.0)
        response.raise_for_status()
    except httpx.HTTPStatusError as e:
        raise RuntimeError(f"Failed to fetch URL {url}: HTTP {e.response.status_code}")
    except httpx.RequestError as e:
        raise RuntimeError(f"Failed to fetch URL {url}: {e}")

    content_type = response.headers.get("content-type", "").lower()

    # JSON endpoint (e.g. dbt manifest.json, API schema)
    if "application/json" in content_type or url.endswith(".json"):
        try:
            data = response.json()
            return json.dumps(data, indent=2)[:80_000], "json"
        except Exception:
            pass

    # YAML endpoint
    if url.endswith((".yaml", ".yml")):
        return response.text[:80_000], "yaml"

    # HTML — strip to content
    if "text/html" in content_type:
        return _strip_html(response.text, url), "html"

    # Plain text fallback
    return response.text[:80_000], "text"


def _strip_html(html: str, url: str) -> str:
    """
    Strip HTML to readable text, with special handling for common doc platforms.
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError(
            "beautifulsoup4 is required for URL content extraction. "
            "Run: pip install beautifulsoup4"
        )

    soup = BeautifulSoup(html, "html.parser")

    # Remove nav, header, footer, sidebar noise
    for tag in soup.find_all(["nav", "header", "footer", "aside", "script", "style", "noscript"]):
        tag.decompose()

    # Platform-specific content selectors
    content = None

    if "confluence" in url:
        content = soup.find(id="main-content") or soup.find(class_="wiki-content")

    elif "notion.so" in url or "notion.site" in url:
        content = soup.find(class_="notion-page-content")

    elif "github.com" in url:
        content = (
            soup.find(id="readme")
            or soup.find(class_="markdown-body")
            or soup.find(id="wiki-body")
        )

    elif "gitbook" in url:
        content = soup.find(class_="page-section")

    elif "developers.hubspot.com" in url or "knowledge.hubspot.com" in url:
        # HubSpot developer/knowledge docs — main article body
        content = (
            soup.find("article")
            or soup.find(class_=re.compile(r"docs-content|article-body|hs-content|developer-docs"))
            or soup.find(id=re.compile(r"hs-content|main-content|docs-content"))
            or soup.find("main")
        )

    elif "stripe.com/docs" in url or "docs.stripe.com" in url:
        content = soup.find(class_=re.compile(r"docs-content|article-body")) or soup.find("main")

    elif "docs.snowflake.com" in url:
        content = soup.find(id="body") or soup.find(class_=re.compile(r"body-content|doc-content")) or soup.find("main")

    if content is None:
        # Generic: find main content block
        content = (
            soup.find("main")
            or soup.find(id="content")
            or soup.find(class_="content")
            or soup.body
        )

    if content is None:
        return soup.get_text(separator="\n", strip=True)[:80_000]

    # Extract text preserving structure
    text = content.get_text(separator="\n", strip=True)
    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text[:80_000]


# ── Structured extraction from raw text ───────────────────────────────────────

def _parse_into_context(
    raw_text: str,
    source_path: str,
    source_type: str,
    document_type: str,
) -> DocumentContext:
    """
    Parse raw extracted text into a structured DocumentContext.

    Strategy depends on source_type:
    - JSON/YAML with known schema (dbt manifest, LookML, OSI spec) → structured parse
    - CSV that looks like a data dictionary → column-header parse
    - Everything else → heuristic text extraction
    """

    notes: list[str] = []
    confidence = "medium"

    # ── Structured formats with known schemas ──────────────────────────────────
    if source_type == "json":
        try:
            data = json.loads(raw_text)
            result = _try_parse_known_json_schema(data, source_path, document_type, notes)
            if result:
                return result
        except json.JSONDecodeError:
            notes.append("Could not parse as JSON — falling back to text extraction")

    if source_type == "yaml":
        try:
            data = yaml.safe_load(raw_text)
            result = _try_parse_known_yaml_schema(data, source_path, document_type, notes)
            if result:
                return result
        except yaml.YAMLError:
            notes.append("Could not parse as YAML — falling back to text extraction")

    if source_type == "csv":
        result = _try_parse_data_dictionary_csv(raw_text, source_path, document_type, notes)
        if result:
            return result

    # ── Heuristic text extraction (PDF, DOCX, HTML, plain text) ───────────────
    tables = _extract_table_defs_from_text(raw_text)
    columns = _extract_column_defs_from_text(raw_text)
    metrics = _extract_metric_defs_from_text(raw_text)
    rules = _extract_business_rules_from_text(raw_text)
    glossary = _extract_glossary_from_text(raw_text)
    join_hints = _extract_join_hints_from_text(raw_text)

    # ── Confidence scoring ─────────────────────────────────────────────────────
    has_structure = any([tables, columns, metrics, rules, glossary, join_hints])
    if not has_structure:
        confidence = "low"
        notes.append(
            "Regex extraction found no structured definitions. "
            "Raw text stored — agents will read it directly (Claude Code style)."
        )
    elif len(columns) < 2 and len(metrics) < 1:
        confidence = "low"
        notes.append("Limited structured content found — raw text also stored for agents.")
    elif len(columns) > 5 or len(metrics) > 3:
        confidence = "high"

    # ── Raw text — always stored for HTML/unstructured docs ────────────────────
    # For developer docs, SaaS schema pages, etc.: store full clean text so agents
    # can read it natively (same approach as Claude Code — no pre-extraction needed).
    # Structured extraction is a bonus when it works; raw text is the fallback.
    raw_summary = None
    if source_type in ("html", "text", "markdown") or confidence in ("low", "medium"):
        raw_summary = raw_text[:50_000]  # 50K chars ≈ ~12K tokens, fits comfortably in context

    # ── LLM extraction — only for low confidence, as an additional enrichment ──
    # Haiku runs on the raw text and adds structured fields when regex can't.
    # The raw text is still stored alongside, so agents have both.
    if confidence == "low":
        llm_result = _llm_extract(raw_text, source_path, document_type, notes)
        if llm_result:
            # Merge raw text into the LLM result so agents always have it
            llm_result.raw_text_summary = raw_summary
            return llm_result

    return DocumentContext(
        source_path=source_path,
        source_type=source_type,
        document_type=document_type,
        extraction_confidence=confidence,
        table_definitions=tables,
        column_definitions=columns,
        metric_definitions=metrics,
        business_rules=rules,
        glossary=glossary,
        join_hints=join_hints,
        raw_text_summary=raw_summary,
        extraction_notes=notes,
    )


# ── Known schema parsers ───────────────────────────────────────────────────────

def _try_parse_known_json_schema(
    data: Any, source_path: str, document_type: str, notes: list[str]
) -> DocumentContext | None:
    """Detect and parse dbt manifest.json, dbt catalog.json, or similar."""

    # dbt manifest.json
    if isinstance(data, dict) and data.get("metadata", {}).get("dbt_schema_version"):
        return _parse_dbt_manifest(data, source_path, notes)

    # dbt catalog.json
    if isinstance(data, dict) and "nodes" in data and "sources" in data and "metadata" in data:
        if "dbt" in str(data.get("metadata", {}).get("generated_at", "")):
            return _parse_dbt_catalog(data, source_path, notes)

    return None


def _try_parse_known_yaml_schema(
    data: Any, source_path: str, document_type: str, notes: list[str]
) -> DocumentContext | None:
    """Detect and parse dbt MetricFlow, dbt schema.yaml, LookML, OSI spec, or generic semantic YAML."""

    if not isinstance(data, dict):
        return None

    # dbt MetricFlow / Semantic Layer (has "semantic_models" key)
    if "semantic_models" in data:
        return _parse_dbt_metricflow_yaml(data, source_path, notes)

    # dbt schema.yaml (has "models" or "sources" but not semantic_models)
    if "models" in data or "sources" in data:
        return _parse_dbt_schema_yaml(data, source_path, notes)

    # LookML-style (has "views" key with dimensions/measures inside)
    if "views" in data or "explores" in data:
        return _parse_lookml_yaml(data, source_path, notes)

    # Generic semantic layer / OSI spec (has "facts" or "dimensions" at top level)
    if "facts" in data or ("dimensions" in data and "metrics" in data):
        return _parse_generic_semantic_yaml(data, source_path, notes)

    return None


def _parse_dbt_manifest(data: dict, source_path: str, notes: list[str]) -> DocumentContext:
    """Parse dbt manifest.json into DocumentContext."""
    tables = []
    columns = []
    metrics = []

    for node_id, node in data.get("nodes", {}).items():
        if node.get("resource_type") not in ("model", "source"):
            continue
        table_name = node.get("name", "")
        desc = node.get("description", "")
        tables.append(TableDefinition(table_name=table_name, description=desc or None))

        for col_name, col_meta in node.get("columns", {}).items():
            columns.append(ColumnDefinition(
                table_name=table_name,
                column_name=col_name,
                description=col_meta.get("description") or None,
                business_name=col_meta.get("meta", {}).get("label") or None,
            ))

    for metric_id, m in data.get("metrics", {}).items():
        metrics.append(MetricDefinition(
            name=m.get("name", metric_id),
            business_name=m.get("label") or m.get("name", metric_id),
            description=m.get("description") or None,
            source_table=m.get("model") or None,
        ))

    notes.append(f"Parsed as dbt manifest. Found {len(tables)} models, {len(metrics)} metrics.")
    return DocumentContext(
        source_path=source_path,
        source_type="json",
        document_type="existing_semantic_layer",
        extraction_confidence="high",
        table_definitions=tables,
        column_definitions=columns,
        metric_definitions=metrics,
        extraction_notes=notes,
    )


def _parse_dbt_catalog(data: dict, source_path: str, notes: list[str]) -> DocumentContext:
    """Parse dbt catalog.json (has physical column type info)."""
    tables = []
    columns = []

    for node_id, node in {**data.get("nodes", {}), **data.get("sources", {})}.items():
        meta = node.get("metadata", {})
        table_name = meta.get("name") or node_id.split(".")[-1]
        tables.append(TableDefinition(table_name=table_name))
        for col_name, col_meta in node.get("columns", {}).items():
            columns.append(ColumnDefinition(
                table_name=table_name,
                column_name=col_name,
                description=col_meta.get("comment") or None,
                data_type_note=col_meta.get("type") or None,
            ))

    notes.append(f"Parsed as dbt catalog. Found {len(tables)} tables.")
    return DocumentContext(
        source_path=source_path,
        source_type="json",
        document_type="existing_semantic_layer",
        extraction_confidence="high",
        table_definitions=tables,
        column_definitions=columns,
        extraction_notes=notes,
    )


def _parse_dbt_schema_yaml(data: dict, source_path: str, notes: list[str]) -> DocumentContext:
    """Parse dbt schema.yaml into DocumentContext."""
    tables = []
    columns = []
    metrics = []

    for model in data.get("models", []) + data.get("sources", []):
        table_name = model.get("name", "")
        tables.append(TableDefinition(
            table_name=table_name,
            description=model.get("description") or None,
        ))
        for col in model.get("columns", []):
            columns.append(ColumnDefinition(
                table_name=table_name,
                column_name=col.get("name", ""),
                description=col.get("description") or None,
                business_name=col.get("meta", {}).get("label") if col.get("meta") else None,
            ))

    for m in data.get("metrics", []):
        metrics.append(MetricDefinition(
            name=m.get("name", ""),
            business_name=m.get("label") or m.get("name", ""),
            description=m.get("description") or None,
            source_table=m.get("model") or None,
            filters=[f.get("field", "") + " = '" + f.get("value", "") + "'"
                     for f in m.get("filters", []) if f.get("field")],
        ))

    notes.append(f"Parsed as dbt schema YAML. {len(tables)} models, {len(metrics)} metrics.")
    return DocumentContext(
        source_path=source_path,
        source_type="yaml",
        document_type="existing_semantic_layer",
        extraction_confidence="high",
        table_definitions=tables,
        column_definitions=columns,
        metric_definitions=metrics,
        extraction_notes=notes,
    )


def _parse_dbt_metricflow_yaml(data: dict, source_path: str, notes: list[str]) -> DocumentContext:
    """
    Parse dbt Semantic Layer / MetricFlow YAML.

    Structure:
      semantic_models:
        - name: orders
          model: ref('orders')
          entities: [{name: order_id, type: primary}]
          dimensions: [{name: order_date, type: time, type_params: {time_granularity: day}}]
          measures: [{name: revenue, agg: sum, expr: amount}]
      metrics:
        - name: revenue
          type: simple
          type_params: {measure: revenue}
    """
    tables = []
    columns = []
    measures = []
    metrics = []

    for sm in data.get("semantic_models", []):
        name = sm.get("name", "")
        model_ref = sm.get("model", "")
        desc = sm.get("description", "")

        # Primary key from entities
        pk = None
        for entity in sm.get("entities", []):
            if entity.get("type") == "primary":
                pk = entity.get("expr") or entity.get("name")
                break

        tables.append(TableDefinition(
            table_name=name,
            description=desc or None,
            table_type="fact",        # semantic_models are fact tables by convention
            primary_key=pk,
            grain_description=sm.get("defaults", {}).get("agg_time_dimension"),
            notes=f"dbt model: {model_ref}" if model_ref else None,
        ))

        # Dimensions → ColumnDefinition with semantic_type
        for dim in sm.get("dimensions", []):
            dim_name = dim.get("name", "")
            dim_type = dim.get("type", "categorical")
            grain = None
            if dim_type == "time":
                sem_type = "time_dimension"
                grain = dim.get("type_params", {}).get("time_granularity")
            else:
                sem_type = "dimension"
            columns.append(ColumnDefinition(
                table_name=name,
                column_name=dim_name,
                description=dim.get("description") or None,
                semantic_type=sem_type,
                time_granularity=grain,
                is_foreign_key=(dim.get("type") == "foreign"),
            ))

        # Entities (foreign keys) → ColumnDefinition
        for entity in sm.get("entities", []):
            etype = entity.get("type", "")
            if etype != "primary":
                columns.append(ColumnDefinition(
                    table_name=name,
                    column_name=entity.get("expr") or entity.get("name", ""),
                    description=entity.get("description") or None,
                    semantic_type="entity",
                    is_foreign_key=(etype == "foreign"),
                ))

        # Measures → MeasureDefinition
        for m in sm.get("measures", []):
            agg = m.get("agg", "SUM").upper()
            measures.append(MeasureDefinition(
                name=m.get("name", ""),
                table_name=name,
                aggregation=agg,
                expression=m.get("expr") or m.get("name"),
                description=m.get("description") or None,
                label=m.get("label") or None,
                non_additive_dimension=m.get("non_additive_dimension", {}).get("name") if isinstance(m.get("non_additive_dimension"), dict) else m.get("non_additive_dimension"),
            ))

    # Metrics block
    for m in data.get("metrics", []):
        mtype = m.get("type", "simple")
        tp = m.get("type_params", {})
        metrics.append(MetricDefinition(
            name=m.get("name", ""),
            business_name=m.get("label") or m.get("name", ""),
            description=m.get("description") or None,
            metric_type=mtype,
            measure_name=tp.get("measure") if mtype == "simple" else None,
            numerator_measure=tp.get("numerator", {}).get("name") if mtype in ("ratio", "derived") else None,
            denominator_measure=tp.get("denominator", {}).get("name") if mtype == "ratio" else None,
            formula=tp.get("expr") if mtype == "derived" else None,
            filters=[f.get("where", "") for f in m.get("filter", []) if f.get("where")],
            dimensions=[d.get("name", "") for d in m.get("dimensions", [])],
        ))

    notes.append(
        f"Parsed as dbt MetricFlow YAML. "
        f"{len(tables)} semantic models, {len(measures)} measures, {len(metrics)} metrics."
    )
    return DocumentContext(
        source_path=source_path,
        source_type="yaml",
        document_type="existing_semantic_layer",
        extraction_confidence="high",
        table_definitions=tables,
        column_definitions=columns,
        measure_definitions=measures,
        metric_definitions=metrics,
        extraction_notes=notes,
    )


def _parse_lookml_yaml(data: dict, source_path: str, notes: list[str]) -> DocumentContext:
    """
    Parse LookML-style YAML (views with dimension/dimension_group/measure fields).

    Structure:
      views:
        - view: orders
          sql_table_name: orders
          dimensions:
            - dimension: order_id
              type: number
              primary_key: yes
            - dimension_group: created
              type: time
              timeframes: [date, week, month]
          measures:
            - measure: count
              type: count
            - measure: total_revenue
              type: sum
              sql: ${amount}
    """
    tables = []
    columns = []
    measures = []

    for view in data.get("views", []):
        view_name = view.get("view", view.get("name", ""))
        sql_table = view.get("sql_table_name", "")
        desc = view.get("description", "")

        pk = None
        for dim in view.get("dimensions", []):
            if str(dim.get("primary_key", "")).lower() in ("yes", "true"):
                pk = dim.get("dimension") or dim.get("name", "")
                break

        tables.append(TableDefinition(
            table_name=view_name,
            description=desc or None,
            primary_key=pk,
            notes=f"sql_table_name: {sql_table}" if sql_table else None,
        ))

        for dim in view.get("dimensions", []):
            dim_name = dim.get("dimension") or dim.get("name", "")
            dim_type = dim.get("type", "string")
            sem_type = "time_dimension" if dim_type == "time" else "dimension"
            timeframes = dim.get("timeframes", [])
            grain = timeframes[0] if timeframes else None
            columns.append(ColumnDefinition(
                table_name=view_name,
                column_name=dim_name,
                description=dim.get("description") or dim.get("label") or None,
                semantic_type=sem_type,
                time_granularity=grain,
            ))

        for m in view.get("measures", []):
            mname = m.get("measure") or m.get("name", "")
            agg = m.get("type", "count").upper()
            measures.append(MeasureDefinition(
                name=mname,
                table_name=view_name,
                aggregation=agg,
                expression=m.get("sql") or mname,
                description=m.get("description") or m.get("label") or None,
                label=m.get("label") or None,
                filters=[m["filters_sql_where"]] if m.get("filters_sql_where") else [],
            ))

    notes.append(
        f"Parsed as LookML YAML. {len(tables)} views, {len(measures)} measures."
    )
    return DocumentContext(
        source_path=source_path,
        source_type="yaml",
        document_type="existing_semantic_layer",
        extraction_confidence="high",
        table_definitions=tables,
        column_definitions=columns,
        measure_definitions=measures,
        extraction_notes=notes,
    )


def _parse_generic_semantic_yaml(data: dict, source_path: str, notes: list[str]) -> DocumentContext:
    """
    Parse generic / OSI-spec semantic layer YAML.

    Handles structures like:
      facts:
        - name: orders
          table: orders
          primary_key: order_id
          grain: one row per order
          measures:
            - name: revenue
              agg: SUM
              expr: amount
          dimensions: [customer_id, order_date]
      dimensions:
        - name: customers
          table: customers
          primary_key: customer_id
          attributes: [country, segment]
      metrics:
        - name: revenue
          measure: revenue
          label: Total Revenue
    """
    tables = []
    columns = []
    measures = []
    metrics = []
    joins = []

    # Fact tables
    for fact in data.get("facts", []):
        tname = fact.get("name") or fact.get("table", "")
        tables.append(TableDefinition(
            table_name=tname,
            description=fact.get("description") or None,
            table_type="fact",
            primary_key=fact.get("primary_key") or fact.get("pk"),
            grain_description=fact.get("grain") or fact.get("grain_description"),
        ))
        for m in fact.get("measures", []):
            measures.append(MeasureDefinition(
                name=m.get("name", ""),
                table_name=tname,
                aggregation=(m.get("agg") or m.get("aggregation", "SUM")).upper(),
                expression=m.get("expr") or m.get("expression") or m.get("name"),
                description=m.get("description") or None,
                label=m.get("label") or None,
                filters=m.get("filters", []) if isinstance(m.get("filters"), list) else [],
            ))
        # Inline dimension references
        for dim_ref in fact.get("dimensions", []):
            if isinstance(dim_ref, str):
                columns.append(ColumnDefinition(
                    table_name=tname,
                    column_name=dim_ref,
                    semantic_type="dimension",
                ))

    # Dimension tables
    for dim_table in data.get("dimensions", []) if isinstance(data.get("dimensions"), list) else []:
        if not isinstance(dim_table, dict):
            continue
        tname = dim_table.get("name") or dim_table.get("table", "")
        tables.append(TableDefinition(
            table_name=tname,
            description=dim_table.get("description") or None,
            table_type="dimension",
            primary_key=dim_table.get("primary_key") or dim_table.get("pk"),
        ))
        for attr in dim_table.get("attributes", []):
            col_name = attr if isinstance(attr, str) else attr.get("name", "")
            col_desc = None if isinstance(attr, str) else attr.get("description")
            col_type = "time_dimension" if (not isinstance(attr, str) and attr.get("type") == "time") else "dimension"
            columns.append(ColumnDefinition(
                table_name=tname,
                column_name=col_name,
                description=col_desc,
                semantic_type=col_type,
            ))

    # Top-level metrics block
    for m in data.get("metrics", []):
        mtype = m.get("type") or ("ratio" if m.get("numerator") else "simple")
        metrics.append(MetricDefinition(
            name=m.get("name", ""),
            business_name=m.get("label") or m.get("name", ""),
            description=m.get("description") or None,
            metric_type=mtype,
            measure_name=m.get("measure") or m.get("measure_name"),
            numerator_measure=m.get("numerator"),
            denominator_measure=m.get("denominator"),
            formula=m.get("formula") or m.get("expr"),
            filters=m.get("filters", []) if isinstance(m.get("filters"), list) else [],
            domain=m.get("domain") or m.get("team"),
        ))

    # Join hints
    for join in data.get("joins", []):
        joins.append(JoinHint(
            left_table=join.get("left") or join.get("from", ""),
            right_table=join.get("right") or join.get("to", ""),
            join_key=join.get("on") or join.get("key", ""),
            cardinality=join.get("cardinality") or join.get("type"),
            description=join.get("description"),
        ))

    notes.append(
        f"Parsed as generic semantic layer YAML. "
        f"{len(tables)} tables ({sum(1 for t in tables if t.table_type == 'fact')} facts, "
        f"{sum(1 for t in tables if t.table_type == 'dimension')} dimensions), "
        f"{len(measures)} measures, {len(metrics)} metrics."
    )
    return DocumentContext(
        source_path=source_path,
        source_type="yaml",
        document_type="existing_semantic_layer",
        extraction_confidence="high",
        table_definitions=tables,
        column_definitions=columns,
        measure_definitions=measures,
        metric_definitions=metrics,
        join_hints=joins,
        extraction_notes=notes,
    )


def _try_parse_data_dictionary_csv(
    raw_text: str, source_path: str, document_type: str, notes: list[str]
) -> DocumentContext | None:
    """
    Try to parse a CSV as a data dictionary.
    Common formats:
      table, column, description, type, ...
      field_name, definition, ...
    """
    reader = csv.DictReader(io.StringIO(raw_text))
    headers = [h.strip().lower() for h in (reader.fieldnames or [])]

    if not headers:
        return None

    # Detect data dictionary columns
    has_table = any(h in headers for h in ["table", "table_name", "model", "source"])
    has_col = any(h in headers for h in ["column", "column_name", "field", "field_name", "attribute"])
    has_desc = any(h in headers for h in ["description", "definition", "desc", "meaning", "notes"])

    if not (has_col and has_desc):
        return None

    columns = []
    for row in reader:
        row_lower = {k.strip().lower(): v for k, v in row.items()}
        table_name = (
            row_lower.get("table") or row_lower.get("table_name") or
            row_lower.get("model") or row_lower.get("source") or "unknown"
        )
        col_name = (
            row_lower.get("column") or row_lower.get("column_name") or
            row_lower.get("field") or row_lower.get("field_name") or
            row_lower.get("attribute") or ""
        )
        description = (
            row_lower.get("description") or row_lower.get("definition") or
            row_lower.get("desc") or row_lower.get("meaning") or
            row_lower.get("notes") or None
        )
        business_name = row_lower.get("business_name") or row_lower.get("label") or None
        is_pii = str(row_lower.get("pii", "")).lower() in ("yes", "true", "1", "y")

        if col_name:
            columns.append(ColumnDefinition(
                table_name=table_name,
                column_name=col_name,
                description=description,
                business_name=business_name,
                is_pii=is_pii,
            ))

    if not columns:
        return None

    notes.append(f"Parsed as data dictionary CSV. Extracted {len(columns)} column definitions.")
    # Derive tables from unique table names
    seen_tables = {}
    for c in columns:
        if c.table_name not in seen_tables:
            seen_tables[c.table_name] = TableDefinition(table_name=c.table_name)

    return DocumentContext(
        source_path=source_path,
        source_type="csv",
        document_type=document_type or "data_dictionary",
        extraction_confidence="high",
        table_definitions=list(seen_tables.values()),
        column_definitions=columns,
        extraction_notes=notes,
    )


# ── LLM-powered extraction ────────────────────────────────────────────────────

def _llm_extract(
    raw_text: str,
    source_path: str,
    document_type: str,
    notes: list[str],
) -> DocumentContext | None:
    """
    Use Claude Haiku to extract structured schema info from unstructured text.

    Activated automatically when regex heuristics yield low confidence.
    Requires ANTHROPIC_API_KEY environment variable. Fails gracefully if absent.
    """
    import os
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        notes.append("ANTHROPIC_API_KEY not set — skipping LLM extraction fallback.")
        return None

    try:
        import anthropic
    except ImportError:
        notes.append("anthropic package not installed — skipping LLM extraction fallback.")
        return None

    # Cap text to keep prompt costs low (Haiku is cheap but let's be reasonable)
    text_snippet = raw_text[:15_000]

    prompt = f"""You are a data architect analyzing documentation about a data warehouse, data source, or semantic layer.
Source URL/path: {source_path}
Document type: {document_type}

Extract ALL of the following from the text below. Return empty arrays/objects where nothing is found.

1. tables — database tables, CRM objects, entities, or data models.
   - name: table/object name
   - description: what it represents
   - table_type: one of "fact", "dimension", "slowly_changing_dimension", "bridge", "staging", "unknown"
     * fact = contains events/transactions with numeric measures (e.g. orders, sessions, payments)
     * dimension = descriptive context (e.g. customers, products, dates, geography)
     * "unknown" if unclear
   - primary_key: the primary key column name if mentioned, else null

2. columns — fields, properties, or attributes.
   - table: table name this belongs to (use "unknown" if unclear)
   - name: column/field name
   - description: what it represents
   - business_name: human-friendly label if different from name, else null
   - data_type: SQL type if mentioned (VARCHAR, INT, TIMESTAMP, etc.), else null
   - semantic_type: one of "dimension", "time_dimension", "entity", "measure", null
     * dimension = categorical attribute used for grouping/filtering
     * time_dimension = a date/timestamp column used for time-based analysis
     * entity = a foreign key or identifier linking to another object/table
     * measure = a raw numeric column that gets aggregated (e.g. amount, quantity, duration)
   - time_granularity: for time_dimensions only — "day", "week", "month", "quarter", "year", else null

3. measures — pre-defined aggregations (SUM, COUNT, AVG, etc.) defined in the semantic model.
   Include these separately from raw columns.
   - name: measure name
   - table: which table/model this measure is defined on
   - aggregation: SUM, COUNT, COUNT_DISTINCT, AVG, MIN, MAX, MEDIAN
   - expression: the underlying column or SQL expression being aggregated, else null
   - description: what it represents, else null
   - label: human-friendly display name, else null

4. metrics — business-level KPIs or calculated measures (often composed from measures).
   - name: metric name
   - business_name: human label, else use name
   - description: what it measures
   - formula: SQL or expression, else null
   - metric_type: one of "simple", "derived", "ratio", "cumulative", "conversion", null
     * simple = directly from a single measure (e.g. total_revenue = SUM(amount))
     * ratio = numerator / denominator (e.g. conversion_rate = conversions / sessions)
     * derived = arithmetic combination of other metrics
     * cumulative = running total over time
   - measure_name: for simple metrics, the underlying measure name, else null
   - numerator_measure: for ratio metrics, the numerator measure name, else null
   - denominator_measure: for ratio metrics, the denominator measure name, else null
   - dimensions: list of dimension names this metric can be sliced by (empty list if unknown)

5. business_rules — filters, conditions, or important notes about how data behaves
   (e.g. "deleted contacts are soft-deleted with deleted_at timestamp",
    "deals in stage Closed Won count toward revenue",
    "sessions under 10 seconds are excluded from engagement metrics").

6. glossary — key business or domain terms with definitions.

7. joins — relationships between tables/objects.
   - left_table, right_table, join_key, cardinality (one_to_one, one_to_many, many_to_many)

Return ONLY valid JSON in exactly this structure (no markdown, no explanation):
{{
  "tables": [{{
    "name": "...", "description": "...", "table_type": "fact|dimension|unknown",
    "primary_key": null
  }}],
  "columns": [{{
    "table": "...", "name": "...", "description": "...", "business_name": null,
    "data_type": null, "semantic_type": "dimension|time_dimension|entity|measure|null",
    "time_granularity": null
  }}],
  "measures": [{{
    "name": "...", "table": "...", "aggregation": "SUM|COUNT|COUNT_DISTINCT|AVG|MIN|MAX|MEDIAN",
    "expression": null, "description": null, "label": null
  }}],
  "metrics": [{{
    "name": "...", "business_name": "...", "description": "...", "formula": null,
    "metric_type": null, "measure_name": null,
    "numerator_measure": null, "denominator_measure": null,
    "dimensions": []
  }}],
  "business_rules": ["..."],
  "glossary": [{{"term": "...", "definition": "..."}}],
  "joins": [{{"left_table": "...", "right_table": "...", "join_key": "...", "cardinality": "one_to_many"}}]
}}

Documentation text:
{text_snippet}"""

    try:
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=6000,
            messages=[{"role": "user", "content": prompt}],
        )
        raw_json = response.content[0].text.strip()

        # Strip markdown fences if the model wraps in ```json
        if raw_json.startswith("```"):
            raw_json = re.sub(r"^```(?:json)?\n?", "", raw_json)
            raw_json = re.sub(r"\n?```$", "", raw_json).strip()

        data = json.loads(raw_json)

        tables = [
            TableDefinition(
                table_name=t["name"],
                description=t.get("description") or None,
                table_type=t.get("table_type") or None,
                primary_key=t.get("primary_key") or None,
            )
            for t in data.get("tables", []) if t.get("name")
        ]
        columns = [
            ColumnDefinition(
                table_name=c.get("table") or "unknown",
                column_name=c["name"],
                description=c.get("description") or None,
                business_name=c.get("business_name") or None,
                data_type_note=c.get("data_type") or None,   # LLM returns "data_type", maps to data_type_note
                semantic_type=c.get("semantic_type") or None,
                time_granularity=c.get("time_granularity") or None,
            )
            for c in data.get("columns", []) if c.get("name")
        ]
        measures = [
            MeasureDefinition(
                name=ms["name"],
                table_name=ms.get("table") or "unknown",
                aggregation=ms.get("aggregation") or "COUNT",
                expression=ms.get("expression") or None,
                description=ms.get("description") or None,
                label=ms.get("label") or None,
            )
            for ms in data.get("measures", []) if ms.get("name")
        ]
        metrics = [
            MetricDefinition(
                name=m["name"],
                business_name=m.get("business_name") or m["name"],
                description=m.get("description") or None,
                formula=m.get("formula") or None,
                metric_type=m.get("metric_type") or None,
                measure_name=m.get("measure_name") or None,
                numerator_measure=m.get("numerator_measure") or None,
                denominator_measure=m.get("denominator_measure") or None,
                dimensions=m.get("dimensions") or [],
            )
            for m in data.get("metrics", []) if m.get("name")
        ]
        rules = [BusinessRule(rule=r) for r in data.get("business_rules", []) if r]
        glossary = [
            GlossaryTerm(term=g["term"], definition=g["definition"])
            for g in data.get("glossary", []) if g.get("term") and g.get("definition")
        ]
        join_hints = [
            JoinHint(
                left_table=j["left_table"],
                right_table=j["right_table"],
                join_key=j.get("join_key") or "",
                cardinality=j.get("cardinality") or None,
            )
            for j in data.get("joins", []) if j.get("left_table") and j.get("right_table")
        ]

        has_structure = any([tables, columns, measures, metrics, rules, glossary, join_hints])
        if not has_structure:
            notes.append("LLM extraction completed but found no structured content in this document.")
            return None

        # Richer confidence signal: semantic layer docs will have measures/metrics
        semantic_richness = len(measures) + len(metrics)
        confidence = (
            "high" if (len(columns) > 5 or len(tables) > 2 or semantic_richness > 2)
            else "medium"
        )
        notes.append(
            f"LLM-assisted extraction (Haiku): {len(tables)} tables, {len(columns)} columns, "
            f"{len(measures)} measures, {len(metrics)} metrics, "
            f"{len(join_hints)} joins, {len(glossary)} glossary terms."
        )

        return DocumentContext(
            source_path=source_path,
            source_type="html",
            document_type=document_type,
            extraction_confidence=confidence,
            table_definitions=tables,
            column_definitions=columns,
            measure_definitions=measures,
            metric_definitions=metrics,
            business_rules=rules,
            glossary=glossary,
            join_hints=join_hints,
            # raw_text_summary will be merged in by _parse_into_context caller
            extraction_notes=notes,
        )

    except Exception as e:
        notes.append(f"LLM extraction failed ({e}) — raw text still available for agents.")
        return None


# ── Heuristic text extraction ─────────────────────────────────────────────────
# These are best-effort pattern matchers for unstructured documents (PDF, DOCX, HTML).
# They are imperfect and agents should treat their output as "hints".

_TABLE_PATTERNS = [
    re.compile(r"(?:table|model|entity)\s*[:\-–—]\s*`?([a-zA-Z_][a-zA-Z0-9_]*)`?", re.IGNORECASE),
    re.compile(r"^#+\s+([a-zA-Z_][a-zA-Z0-9_]+)\s*$", re.MULTILINE),
]

_COLUMN_PATTERNS = [
    # "column_name (type): description"
    re.compile(
        r"`([a-zA-Z_][a-zA-Z0-9_]*)`\s*(?:\(([^)]+)\))?\s*[:\-–—]\s*(.{10,200})",
        re.IGNORECASE
    ),
    # "• field_name — description"
    re.compile(
        r"[•\-\*]\s+([a-zA-Z_][a-zA-Z0-9_]*)\s*[:\-–—]\s*(.{10,200})",
        re.IGNORECASE
    ),
]

_METRIC_PATTERNS = [
    re.compile(
        r"(?:metric|measure|kpi|indicator)\s*[:\-–—]\s*([^\n]+)\n\s*(?:formula|calculation|definition)\s*[:\-–—]\s*([^\n]+)",
        re.IGNORECASE
    ),
]

_RULE_PATTERNS = [
    re.compile(
        r"(?:business rule|rule|logic|note|important|filter)\s*[:\-–—]\s*(.{20,400})",
        re.IGNORECASE
    ),
]

_GLOSSARY_PATTERNS = [
    # "**Term**: definition" (markdown)
    re.compile(r"\*\*([^*]+)\*\*\s*[:\-–—]\s*(.{20,400})", re.IGNORECASE),
    # "Term (abbreviation): definition"
    re.compile(r"^([A-Z][a-zA-Z\s]+(?:\([A-Z]+\))?)\s*[:\-–—]\s*(.{20,400})", re.MULTILINE),
]

_JOIN_PATTERNS = [
    re.compile(
        r"([a-zA-Z_][a-zA-Z0-9_]*)\s+(?:joins?|links?|relates?\s+to|foreign\s+key(?:\s+to)?)\s+([a-zA-Z_][a-zA-Z0-9_]*)\s+(?:on|via|using)?\s*`?([a-zA-Z_][a-zA-Z0-9_]*)`?",
        re.IGNORECASE
    ),
]


def _extract_table_defs_from_text(text: str) -> list[TableDefinition]:
    seen = set()
    tables = []
    for pat in _TABLE_PATTERNS:
        for m in pat.finditer(text):
            name = m.group(1).strip()
            if name.lower() not in seen and len(name) > 2:
                seen.add(name.lower())
                tables.append(TableDefinition(table_name=name))
    return tables[:50]  # cap


def _extract_column_defs_from_text(text: str) -> list[ColumnDefinition]:
    seen = set()
    columns = []
    for pat in _COLUMN_PATTERNS:
        for m in pat.finditer(text):
            groups = m.groups()
            col_name = groups[0].strip()
            description = groups[-1].strip()[:500]
            if col_name.lower() not in seen and len(col_name) > 1:
                seen.add(col_name.lower())
                columns.append(ColumnDefinition(
                    table_name="unknown",  # table attribution requires more context
                    column_name=col_name,
                    description=description,
                ))
    return columns[:200]  # cap


def _extract_metric_defs_from_text(text: str) -> list[MetricDefinition]:
    metrics = []
    for pat in _METRIC_PATTERNS:
        for m in pat.finditer(text):
            name = m.group(1).strip()
            formula = m.group(2).strip() if len(m.groups()) > 1 else None
            metrics.append(MetricDefinition(
                name=re.sub(r"\s+", "_", name.lower()),
                business_name=name,
                formula=formula,
            ))
    return metrics[:50]


def _extract_business_rules_from_text(text: str) -> list[BusinessRule]:
    rules = []
    seen = set()
    for pat in _RULE_PATTERNS:
        for m in pat.finditer(text):
            rule_text = m.group(1).strip()[:400]
            if rule_text not in seen and len(rule_text) > 20:
                seen.add(rule_text)
                rules.append(BusinessRule(rule=rule_text))
    return rules[:30]


def _extract_glossary_from_text(text: str) -> list[GlossaryTerm]:
    terms = []
    seen = set()
    for pat in _GLOSSARY_PATTERNS:
        for m in pat.finditer(text):
            term = m.group(1).strip()
            definition = m.group(2).strip()[:400]
            if term.lower() not in seen and 2 < len(term) < 60:
                seen.add(term.lower())
                terms.append(GlossaryTerm(term=term, definition=definition))
    return terms[:100]


def _extract_join_hints_from_text(text: str) -> list[JoinHint]:
    hints = []
    for pat in _JOIN_PATTERNS:
        for m in pat.finditer(text):
            hints.append(JoinHint(
                left_table=m.group(1),
                right_table=m.group(2),
                join_key=m.group(3) if len(m.groups()) > 2 else "",
            ))
    return hints[:50]
